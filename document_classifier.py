import os
import pdfplumber
import docx
import logging

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# ==============================================================================
# Extract Text
# ==============================================================================

def extract_text_from_file(file_path):
    text = ""
    if not os.path.exists(file_path):
        print(f"Error: File not found at '{file_path}'")
        return None

    if file_path.endswith('.pdf'):
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text
        except Exception as e:
            print(f"Could not read PDF file '{file_path}': {e}")
            return None
    elif file_path.endswith('.docx'):
        try:
            doc = docx.Document(file_path)
            # Extract text from paragraphs
            for para in doc.paragraphs:
                text += para.text + "\n"
            
            # Extract text from tables (important for resumes)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
        except Exception as e:
            print(f"Could not read DOCX file '{file_path}': {e}")
            return None
    elif file_path.endswith('.txt'):
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
        except Exception as e:
            print(f"Could not read TXT file '{file_path}': {e}")
            return None
    else:
        print(f"Unsupported file type: '{file_path}'")
        return None
        
    return text.lower()

# ==============================================================================
# (KEYWORD-BASED)
# ==============================================================================

def classify_document(text):

    try: 
        if not text or not text.strip():
            return "Unclassified (Empty)"

        # Enhanced classification with weighted keywords
        classification_rules = {
            "Resume / CV": [
                # Strong indicators (weight: 3)
                ("curriculum vitae", 3), ("resume", 3), ("cv", 2),
                # Contact & personal info (weight: 2)
                ("email:", 2), ("phone:", 2), ("address:", 2), ("linkedin", 2),
                ("portfolio", 2), ("github", 2),
                # Education section (weight: 2)
                ("education", 2), ("bachelor", 2), ("master", 2), ("degree", 2),
                ("university", 2), ("college", 2), ("graduated", 2), ("gpa", 2),
                # Work experience (weight: 2)
                ("work experience", 2), ("professional experience", 2), 
                ("employment history", 2), ("job title", 1), ("position", 1),
                # Skills & abilities (weight: 1.5)
                ("skills", 1.5), ("technical skills", 2), ("competencies", 1.5),
                ("proficient", 1.5), ("expertise", 1.5), ("abilities", 1.5),
                # Career objective (weight: 1.5)
                ("objective", 1.5), ("career objective", 2), ("professional summary", 2),
                ("summary", 1), ("profile", 1),
                # Other sections (weight: 1)
                ("references", 1), ("certifications", 1.5), ("achievements", 1.5),
                ("accomplishments", 1.5), ("projects", 1), ("volunteer", 1),
                ("awards", 1.5), ("honors", 1.5), ("languages", 1),
                ("qualifications", 1.5), ("training", 1)
            ],
            "Offer Letter / Employment Contract": [
                # Strong indicators (weight: 3)
                ("offer of employment", 3), ("employment contract", 3), 
                ("offer letter", 3), ("employment agreement", 3),
                # Contract terms (weight: 2)
                ("start date", 2), ("salary", 2), ("compensation", 2),
                ("terms and conditions", 2), ("probationary period", 2),
                ("benefits", 2), ("position", 1.5), ("job title", 1.5),
                # Legal terms (weight: 1.5)
                ("at-will employment", 2), ("employee handbook", 1.5),
                ("confidentiality", 1.5), ("non-compete", 1.5),
                ("acceptance", 1.5), ("contingent upon", 1.5),
                # Other terms (weight: 1)
                ("reporting to", 1), ("department", 1), ("base salary", 2),
                ("annual salary", 2), ("hourly rate", 2)
            ],
            "Resignation Letter": [
                # Strong indicators (weight: 3)
                ("resign", 3), ("resignation", 3), ("letter of resignation", 3),
                # Notice period (weight: 2)
                ("two weeks notice", 2), ("notice period", 2), ("last day", 2),
                ("final day", 2), ("effective date", 2),
                # Departure phrases (weight: 2)
                ("leaving the company", 2), ("step down", 2), ("moving on", 1.5),
                ("terminate my employment", 2), ("end my employment", 2),
                # Other phrases (weight: 1)
                ("thank you for the opportunity", 1), ("grateful", 1),
                ("transition", 1), ("handover", 1)
            ],
            "Leave Request Form": [
                # Strong indicators (weight: 3)
                ("leave request", 3), ("leave application", 3), ("request for leave", 3),
                # Leave types (weight: 2)
                ("sick leave", 2), ("vacation", 2), ("annual leave", 2),
                ("personal leave", 2), ("maternity leave", 2), ("paternity leave", 2),
                ("medical leave", 2), ("emergency leave", 2),
                # Request phrases (weight: 2)
                ("time off", 2), ("leave of absence", 2), ("dates of leave", 2),
                ("requesting leave", 2), ("apply for leave", 2),
                # Other terms (weight: 1)
                ("from date", 1), ("to date", 1), ("duration", 1),
                ("reason for leave", 1.5), ("approval", 1)
            ],
            "Employee Warning / Disciplinary Action": [
                # Strong indicators (weight: 3)
                ("disciplinary action", 3), ("written warning", 3), 
                ("disciplinary notice", 3), ("formal warning", 3),
                # Violation terms (weight: 2)
                ("misconduct", 2), ("violation", 2), ("policy violation", 2),
                ("breach", 2), ("infraction", 2),
                # Improvement terms (weight: 2)
                ("performance improvement", 2), ("corrective action", 2),
                ("improvement plan", 2), ("consequences", 2),
                # Other terms (weight: 1.5)
                ("reprimand", 1.5), ("behavioral issues", 1.5),
                ("unacceptable", 1.5), ("termination", 1.5),
                ("suspension", 1.5), ("probation", 1)
            ],
            "Performance Review": [
                # Strong indicators (weight: 3)
                ("performance review", 3), ("performance appraisal", 3),
                ("performance evaluation", 3), ("annual review", 3),
                # Evaluation terms (weight: 2)
                ("evaluation", 2), ("assessment", 2), ("self-assessment", 2),
                ("rating", 2), ("score", 1.5),
                # Goals & objectives (weight: 2)
                ("goals", 2), ("objectives", 2), ("performance goals", 2),
                ("key performance indicators", 2), ("kpi", 2), ("targets", 2),
                # Feedback terms (weight: 1.5)
                ("feedback", 1.5), ("strengths and weaknesses", 2),
                ("areas for improvement", 1.5), ("achievements", 1.5),
                ("accomplishments", 1.5),
                # Career development (weight: 1)
                ("career development", 1.5), ("development plan", 1.5),
                ("growth", 1), ("promotion", 1)
            ]
        }

        # Calculate weighted scores
        scores = {label: 0.0 for label in classification_rules}
        match_counts = {label: 0 for label in classification_rules}
        
        for label, weighted_keywords in classification_rules.items():
            for keyword, weight in weighted_keywords:
                if keyword in text:
                    scores[label] += weight
                    match_counts[label] += 1
        
        # Find the best match
        highest_score = 0.0
        best_label = "Unclassified"
        
        for label, score in scores.items():
            if score > highest_score:
                highest_score = score
                best_label = label
        
        # Require a minimum threshold to avoid false positives
        # At least 2 matches or a weighted score of 3.0
        if highest_score < 3.0 and match_counts[best_label] < 2:
            return "Unclassified"
        
        logger.info(f"Classification: {best_label} (score: {highest_score:.1f}, matches: {match_counts[best_label]})")
        return best_label
    
    except Exception as e:
        logger.error(f"Classification error: {str(e)}")
        return "Unclassified (Error)"

# ==============================================================================
# MAIN EXECUTION
# ==============================================================================

if __name__ == "__main__":
    
    # file path
    file_to_test = "C:/Users/PLPASIG/Downloads/DMS/OCR file type/dummy-R.docx" 
    
    print("--- Document Classification System ---")
    print(f"\nProcessing file: '{file_to_test}'")
    
    extracted_text = extract_text_from_file(file_to_test)
    
    if extracted_text:
        document_label = classify_document(extracted_text)
        print("\n" + "="*40)
        print(f"✅  Result: This document is classified as: {document_label}")
        print("="*40)
    else:
        print("\n" + "="*40)
        print("❌ Result: Could not process the file.")
        print("="*40)

    print("\n--- Process Finished ---")