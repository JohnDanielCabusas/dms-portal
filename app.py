from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
import os
import tempfile
from dotenv import load_dotenv
from document_classifier import extract_text_from_file, classify_document
from database.database_config import DatabaseConfig
import json
from datetime import datetime
import hashlib
import secrets
from werkzeug.utils import secure_filename
import requests 
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from docx import Document as DocxDocument
import pdfplumber
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
import pandas as pd
from html import escape
from cryptography.fernet import Fernet, InvalidToken

# Load environment variables from .env file
load_dotenv()


def _init_fernet():
    """
    Initialize a single Fernet instance for encrypting sensitive values.
    Uses the DMS_FERNET_KEY environment variable when available,
    otherwise falls back to None (and callers should degrade gracefully).
    """
    key = os.environ.get("DMS_FERNET_KEY")
    if not key:
        print("WARNING: DMS_FERNET_KEY not found in environment. Fernet encryption disabled. Using hash-only mode.")
        return None
    try:
        # Fernet keys must be 32 bytes, base64-encoded (44 characters)
        # If the key is already a base64 string, use it directly
        if isinstance(key, str):
            # Remove any whitespace
            key = key.strip()
            # If it's not already bytes, try to decode from base64
            try:
                # Fernet expects base64-encoded key (44 chars)
                if len(key) == 44:
                    # This is likely a base64-encoded key
                    key_bytes = key.encode("utf-8")
                else:
                    # Try to decode if it's a hex string or other format
                    key_bytes = key.encode("utf-8")
            except:
                key_bytes = key.encode("utf-8")
        else:
            key_bytes = key
        
        fernet_instance = Fernet(key_bytes)
        print("SUCCESS: Fernet encryption initialized successfully.")
        return fernet_instance
    except Exception as e:
        print(f"WARNING: Invalid DMS_FERNET_KEY, falling back to hash-only mode: {e}")
        print(f"Key format check: length={len(key) if key else 0}, type={type(key)}")
        return None


FERNET = _init_fernet()


def current_encryption_version():
    """Return a short tag describing the mechanism protecting file paths."""
    return 'fernet_v1' if FERNET else 'hash_v1'


def encrypt_sensitive(value: str) -> str:
    """
    Encrypt a sensitive string value using Fernet if configured,
    otherwise fall back to a one-way SHA-256 hash.
    """
    if value is None:
        return None

    # Ensure string
    if not isinstance(value, str):
        value = str(value)

    if FERNET is None:
        # Backwards-compatible behaviour: deterministic hash
        return hashlib.sha256(value.encode("utf-8")).hexdigest()

    token = FERNET.encrypt(value.encode("utf-8"))
    return token.decode("utf-8")


def decrypt_sensitive(token: str) -> str | None:
    """
    Decrypt a Fernet-encrypted token. Returns None if decryption
    is not possible (no key configured or invalid token).
    """
    if token is None or FERNET is None:
        return None

    if not isinstance(token, str):
        token = str(token)

    try:
        value = FERNET.decrypt(token.encode("utf-8"))
        return value.decode("utf-8")
    except InvalidToken:
        return None
    except Exception as e:
        print(f"WARNING: decrypt_sensitive failed: {e}")
        return None


def hash_file_path(file_path):
    """
    Backwards-compatible helper for protecting file paths before storing
    them in the database. Now uses Fernet when configured for reversible
    encryption, and otherwise falls back to a one-way SHA-256 hash.
    """
    if not file_path:
        return None
    return encrypt_sensitive(file_path)

app = Flask(__name__)
CORS(app)

# Define UPLOAD_FOLDER outside your code directory
UPLOAD_FOLDER = os.path.join(tempfile.gettempdir(), 'dms_uploads')  # Or use a custom path like '/var/dms/uploads'
ALLOWED_EXTENSIONS = {
    'txt', 'pdf',
    'png', 'jpg', 'jpeg', 'gif', 'bmp',
    'doc', 'docx',
    'xls', 'xlsx',
    'ppt', 'pptx',
    'zip'
}

# Add configuration for file storage
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024  # 50MB max file size

# Now check and create the upload folder
if not os.path.exists(UPLOAD_FOLDER):
    os.makedirs(UPLOAD_FOLDER)

def allowed_file(filename):
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def normalize_profile_image_token(value):
    if not value:
        return None
    value = value.strip()
    if not value:
        return None
    return os.path.basename(value)

def build_profile_image_url(value):
    if not value:
        return None
    if value.startswith('http'):
        return value
    if value.startswith('/'):
        return value
    return f"/api/uploads/profile_images/{value}"

def serialize_user_record(user):
    if not user:
        return user
    user_copy = dict(user)
    user_copy['profile_image'] = build_profile_image_url(user_copy.get('profile_image'))
    return user_copy


def build_company_logo_url(token):
    if not token:
        return None
    if token.startswith('http'):
        return token
    if token.startswith('/'):
        return token
    return f"/api/uploads/company_logo/{token}"


def serialize_settings_record(settings):
    if not settings:
        return settings
    s = dict(settings)
    s['company_logo_url'] = build_company_logo_url(s.get('company_logo'))
    return s

def save_uploaded_file(file, user_id, custom_filename=None, allow_overwrite=False):
    try:
        if not file or file.filename == '':
            return None, None, None, "No file selected"
            
        if not allowed_file(file.filename):
            return None, None, None, "File type not allowed"
        
        # Keep original filename but make it safe
        original_filename = secure_filename(file.filename)
        
        # Use custom filename if provided, otherwise use original
        if custom_filename:
            # Ensure custom filename has extension
            if '.' not in custom_filename:
                extension = original_filename.rsplit('.', 1)[1].lower() if '.' in original_filename else ''
                if extension:
                    custom_filename = f"{custom_filename}.{extension}"
            target_filename = secure_filename(custom_filename)
        else:
            target_filename = original_filename
        
        # Create organized folder structure to avoid conflicts
        today = datetime.now()
        year = today.year
        month = today.month
        user_folder = f"user_{user_id}"
        
        # Create folder path: uploads/2024/03/user_1/
        folder_path = os.path.join(app.config['UPLOAD_FOLDER'], str(year), str(month).zfill(2), user_folder)
        os.makedirs(folder_path, exist_ok=True)
        
        # Get only ACTIVE files from database to check for name conflicts
        query = "SELECT name, file_id FROM files WHERE user_id = %s AND status = 'active' AND name = %s"
        connection = DatabaseConfig.get_connection()
        cursor = connection.cursor(dictionary=True)
        cursor.execute(query, (user_id, target_filename))
        existing_file = cursor.fetchone()
        cursor.close()
        connection.close()
        
        final_filename = target_filename
        file_path = os.path.join(folder_path, final_filename)
        
        # If file exists and overwrite is allowed, delete the existing file
        if existing_file and allow_overwrite:
            # Delete existing file record and physical file
            existing_file_id = existing_file['file_id']
            existing_file_record = DMSDatabase.get_file_by_id(existing_file_id)
            if existing_file_record:
                old_file_path = existing_file_record.get('file_path')
                if old_file_path and os.path.exists(old_file_path):
                    try:
                        os.remove(old_file_path)
                    except Exception as e:
                        print(f"Warning: Could not delete old file: {e}")
                # Delete the old file record
                DMSDatabase.delete_file(existing_file_id)
        elif existing_file and not allow_overwrite:
            # Auto-number if file exists and overwrite not allowed
            base_name = target_filename.rsplit('.', 1)[0] if '.' in target_filename else target_filename
            extension = target_filename.rsplit('.', 1)[1].lower() if '.' in target_filename else ''
            
            # Check if base_name already has a number pattern like "file (1)" or "file (2)"
            import re
            number_pattern = r'\s*\(\d+\)\s*$'
            base_name_clean = re.sub(number_pattern, '', base_name)
            
            query = "SELECT name FROM files WHERE user_id = %s AND status = 'active' AND name LIKE %s"
            connection = DatabaseConfig.get_connection()
            cursor = connection.cursor(dictionary=True)
            like_pattern = f"{base_name_clean}%.{extension}" if extension else f"{base_name_clean}%"
            cursor.execute(query, (user_id, like_pattern))
            existing_files = cursor.fetchall()
            cursor.close()
            connection.close()
            
            # Extract existing numbers from filenames
            existing_numbers = []
            for f in existing_files:
                name = f['name']
                if extension:
                    name_without_ext = name.rsplit('.', 1)[0]
                else:
                    name_without_ext = name
                
                # Check if it matches our pattern
                match = re.search(r'\(\d+\)$', name_without_ext)
                if match:
                    num = int(match.group(0)[1:-1])  # Extract number from "(1)"
                    existing_numbers.append(num)
                elif name_without_ext == base_name_clean:
                    existing_numbers.append(0)  # Original file without number
            
            # Find the next available number
            counter = 1
            while counter in existing_numbers:
                counter += 1
            
            if extension:
                final_filename = f"{base_name_clean} ({counter}).{extension}"
            else:
                final_filename = f"{base_name_clean} ({counter})"
            
            file_path = os.path.join(folder_path, final_filename)
        
        # Actually save the file
        file.save(file_path)
        
        # Verify file was saved
        if not os.path.exists(file_path):
            return None, None, None, "Failed to save file to server"
            
        print(f"DEBUG: File successfully saved at: {file_path}")
        file_path_hash = hash_file_path(file_path)
        return final_filename, file_path, file_path_hash, None
        
    except Exception as e:
        print(f"Error saving file: {e}")
        return None, None, None, f"File save error: {str(e)}"

def resolve_file_path(file_record):
    """Return actual file path on disk for a stored file record.
    
    The file_path stored in the database may be:
    - Encrypted with Fernet (if DMS_FERNET_KEY is set)
    - Hashed with SHA-256 (fallback if no Fernet key)
    - Plain path (legacy records)
    
    This function attempts to decrypt the stored path first, then falls back
    to searching by filename if decryption fails or the path doesn't exist.
    """
    if not file_record:
        return None
    
    stored_path = file_record.get('file_path')
    if not stored_path:
        # Fallback: search by filename
        return _resolve_file_path_by_name(file_record)
    
    # Try to decrypt if Fernet is available
    if FERNET is not None:
        decrypted_path = decrypt_sensitive(stored_path)
        if decrypted_path and os.path.exists(decrypted_path):
            return decrypted_path
    
    # If stored_path is a plain path (legacy) and exists, use it
    if os.path.exists(stored_path):
        return stored_path
    
    # Fallback: search by filename
    return _resolve_file_path_by_name(file_record)

def _resolve_file_path_by_name(file_record):
    """Fallback method to find file by searching for its name in upload folder."""
    upload_folder = app.config.get('UPLOAD_FOLDER')
    target_names = [file_record.get('name'), file_record.get('original_name')]
    if not upload_folder or not os.path.exists(upload_folder):
        return None
    
    target_names = [name for name in target_names if name]
    if not target_names:
        return None
    
    for root, dirs, files in os.walk(upload_folder):
        for f in files:
            if f in target_names:
                return os.path.join(root, f)
    return None

def get_file_extension(file_record):
    filename = file_record.get('name') or file_record.get('original_name') or ''
    return os.path.splitext(filename)[1].lower().lstrip('.')

def extract_text_for_editing(file_record, actual_path):
    """
    Convert supported file types into plain text for the inline editor.
    Returns text content or raises ValueError for unsupported formats.
    """
    file_type = (file_record.get('type') or get_file_extension(file_record) or '').lower()
    
    if file_type in ['txt', 'md', 'html', 'css', 'js', 'json', 'xml', 'csv']:
        try:
            with open(actual_path, 'r', encoding='utf-8') as f:
                return f.read()
        except UnicodeDecodeError:
            with open(actual_path, 'r', encoding='latin-1') as f:
                return f.read()
    
    if file_type == 'docx':
        doc = DocxDocument(actual_path)
        return "\n".join(paragraph.text for paragraph in doc.paragraphs)
    
    if file_type == 'pdf':
        with pdfplumber.open(actual_path) as pdf:
            pages = []
            for page in pdf.pages:
                page_text = page.extract_text() or ""
                pages.append(page_text)
        text = "\n".join(pages).strip()
        if not text:
            raise ValueError("Unable to extract text from PDF for editing.")
        return text
    
    raise ValueError(f'File type ".{file_type}" is not supported for inline editing.')

def save_text_back_to_file(file_record, actual_path, content):
    """
    Persist edited text back into the original file type.
    Supports text, DOCX, and PDF generation.
    """
    file_type = (file_record.get('type') or get_file_extension(file_record) or '').lower()
    
    if file_type in ['txt', 'md', 'html', 'css', 'js', 'json', 'xml', 'csv']:
        with open(actual_path, 'w', encoding='utf-8') as f:
            f.write(content)
        return
    
    if file_type == 'docx':
        doc = DocxDocument()
        for line in content.splitlines():
            doc.add_paragraph(line)
        doc.save(actual_path)
        return
    
    if file_type == 'pdf':
        c = canvas.Canvas(actual_path, pagesize=letter)
        width, height = letter
        margin = 72
        y = height - margin
        line_height = 14
        for line in content.splitlines():
            text_line = line[:180]
            c.drawString(margin, y, text_line)
            y -= line_height
            if y <= margin:
                c.showPage()
                y = height - margin
        c.save()
        return
    
    raise ValueError(f'File type ".{file_type}" cannot be saved through the inline editor.')

def get_request_user():
    user_id = request.headers.get('X-User-Id')
    if not user_id:
        return None
    try:
        return DMSDatabase.get_user_by_id(int(user_id))
    except (ValueError, TypeError):
        return None

def merge_permission(current, new):
    order = {'viewer': 1, 'editor': 2, 'owner': 3, 'admin': 4}
    if current is None:
        return new
    return new if order.get(new, 0) > order.get(current, 0) else current

def get_user_file_permission(file_record, user):
    """
    Determine the permission level a user has for a specific file.
    Returns 'owner', 'editor', 'viewer', 'admin', or None if no access.
    """
    if not user or not file_record:
        return None
    
    if user['role'] == 'system_admin':
        return 'admin'
    
    if file_record['user_id'] == user['user_id']:
        return 'owner'
    
    share = DatabaseConfig.execute_query(
        "SELECT permission FROM file_shares WHERE file_id=%s AND shared_with_user_id=%s",
        (file_record['file_id'], user['user_id']),
        fetch_one=True
    )
    if share:
        return share.get('permission', 'viewer')
    
    dept_id = user.get('department_id')
    if dept_id:
        dept_share = DatabaseConfig.execute_query(
            "SELECT permission FROM file_shares WHERE file_id=%s AND shared_with_department_id=%s",
            (file_record['file_id'], dept_id),
            fetch_one=True
        )
        if dept_share:
            return dept_share.get('permission', 'viewer')
    
    # Check if file is in a workspace where user is a member or owner
    workspace_access = DatabaseConfig.execute_query(
        """SELECT fw.workspace_id 
           FROM file_workspaces fw
           JOIN workspaces w ON fw.workspace_id = w.workspace_id
           LEFT JOIN workspace_members wm ON fw.workspace_id = wm.workspace_id AND wm.user_id = %s
           WHERE fw.file_id = %s 
           AND (w.user_id = %s OR wm.user_id IS NOT NULL)""",
        (user['user_id'], file_record['file_id'], user['user_id']),
        fetch_one=True
    )
    if workspace_access:
        return 'viewer'
    
    return None

# Update the upload-file route
@app.route('/api/upload-file', methods=['POST'])
def upload_file():
    try:
        print(f"DEBUG: Upload request received")
        print(f"DEBUG: Files in request: {list(request.files.keys())}")
        
        if 'file' not in request.files:
            return jsonify({'error': 'No file provided'}), 400
        
        file = request.files['file']
        print(f"DEBUG: File received: {file.filename}")
        
        if file.filename == '':
            return jsonify({'error': 'Please select a file to upload'}), 400
        
        user_id = request.form.get('user_id')
        workspace_ids = request.form.getlist('workspace_ids')  # Get multiple workspace IDs
        category_id = request.form.get('category_id')
        department_id = request.form.get('department_id')
        send_email = request.form.get('send_email', 'false').lower() == 'true'
        
        print(f"DEBUG: User ID: {user_id}, Workspace IDs: {workspace_ids}, Category ID: {category_id}, Department ID: {department_id}, Send Email: {send_email}")
        
        if not user_id:
            return jsonify({'error': 'User ID is required'}), 400
        
        # Check file size
        file.seek(0, 2)
        file_size = file.tell()
        file.seek(0)
        
        settings = DMSDatabase.get_settings()
        max_file_size = (settings.get('max_file_mb') or 50) * 1024 * 1024
        
        if file_size > max_file_size:
            max_mb = settings.get('max_file_mb') or 50
            return jsonify({'error': f'File size exceeds maximum allowed size of {max_mb}MB'}), 400
        
        # Check file type
        if not allowed_file(file.filename):
            allowed_types = settings.get('allowed_types', 'pdf,doc,docx,jpg,jpeg,png,txt')
            return jsonify({'error': f'File type not allowed. Allowed types: {allowed_types}'}), 400
        
        # Get custom filename and overwrite flag if provided
        custom_filename = request.form.get('custom_filename')
        allow_overwrite = request.form.get('allow_overwrite', 'false').lower() == 'true'
        
        # Save the file
        final_filename, file_path, file_path_hash, error_msg = save_uploaded_file(file, user_id, custom_filename, allow_overwrite)
        if not file_path:
            return jsonify({'error': error_msg or 'Failed to save file'}), 400

        print(f"DEBUG: File saved successfully: {final_filename} at {file_path}")

        # Document classification logic
        classification_result = {}
        try:
            extracted_text = extract_text_from_file(file_path)
            if extracted_text:
                document_type = classify_document(extracted_text)
                confidence = calculate_confidence(extracted_text, document_type)
                classification_result = {
                    'document_type': document_type,
                    'confidence': confidence,
                    'text_sample': extracted_text[:200] + '...' if len(extracted_text) > 200 else extracted_text
                }
                
                # Auto-create category if it doesn't exist
                if document_type and document_type != "Unclassified":
                    existing_category = DMSDatabase.get_category_by_name(document_type)
                    if not existing_category:
                        DMSDatabase.create_category(
                            name=document_type,
                            description=f"Auto-created category for {document_type}",
                            auto_created=True,
                            is_unclassified=False
                        )
        except Exception as e:
            print(f"Classification error: {e}")
            classification_result = {
                'document_type': 'Unclassified',
                'confidence': 'Low',
                'error': str(e)
            }
        
        # Create file record in database - set workspace_id to NULL initially
        file_data = {
            'name': final_filename,
            'original_name': file.filename,
            'size': os.path.getsize(file_path),
            'type': file.filename.rsplit('.', 1)[1].lower(),
            'user_id': user_id,
            'workspace_id': None,  # Set to NULL - we use file_workspaces table now
            'category_id': category_id if category_id and category_id != 'null' else None,
            'document_type': classification_result.get('document_type'),
            'classification_confidence': classification_result.get('confidence'),
            'classification_error': classification_result.get('error'),
            'text_sample': classification_result.get('text_sample'),
            'shared': False,
            'file_path': file_path_hash,  # Store only the protected path/token in the DB
            'is_encrypted': bool(file_path_hash),
            'encryption_version': current_encryption_version()
        }

        print(f"DEBUG: Creating file record with data: {file_data}")
        print(f"DEBUG: Real file path: {file_path} | Hash stored: {file_path_hash}")
        print(f"DEBUG: Encryption status - is_encrypted: {file_data.get('is_encrypted')}, encryption_version: {file_data.get('encryption_version')}")
        print(f"DEBUG: Fernet initialized: {FERNET is not None}")

        file_id = DMSDatabase.create_file(file_data)
        
        # Send email notification if requested
        if send_email and department_id:
            try:
                # Get department users
                users = DMSDatabase.get_all_users()
                department_users = [u for u in users if u.get('department_id') == int(department_id) and u.get('status') == 'active']
                
                if department_users:
                    # Send email notification
                    current_user = DMSDatabase.get_user_by_id(user_id)
                    uploader_name = current_user.get('name', 'Unknown') if current_user else 'Unknown'
                    
                    # Use existing email sending logic
                    smtp_server = "smtp.gmail.com"
                    smtp_port = 587
                    sender_email = "plp.dmshr@gmail.com"
                    sender_password = "thpn ttsr pxcw zchi"
                    
                    from email.mime.text import MIMEText
                    from email.mime.multipart import MIMEMultipart
                    import smtplib
                    
                    for dept_user in department_users:
                        try:
                            recipient_email = dept_user.get('email')
                            if not recipient_email:
                                continue
                            
                            subject = f"New File Uploaded: {final_filename}"
                            html_content = f"""
                            <!DOCTYPE html>
                            <html>
                            <head>
                                <style>
                                    body {{ font-family: Arial, sans-serif; background-color: #f4f4f4; padding: 20px; }}
                                    .container {{ max-width: 600px; margin: 0 auto; background: white; padding: 30px; border-radius: 10px; box-shadow: 0 2px 10px rgba(0,0,0,0.1); }}
                                    .header {{ text-align: center; margin-bottom: 30px; color: #3b82f6; }}
                                    .content {{ margin: 20px 0; }}
                                    .info-box {{ background: #f8fafc; padding: 15px; border-radius: 8px; margin: 15px 0; }}
                                    .footer {{ margin-top: 30px; text-align: center; color: #64748b; font-size: 14px; }}
                                </style>
                            </head>
                            <body>
                                <div class='container'>
                                    <div class='header'>
                                        <h2>New File Uploaded</h2>
                                    </div>
                                    <div class='content'>
                                        <p>Hello {dept_user.get('name', 'User')},</p>
                                        <p>A new file has been uploaded to your department:</p>
                                        <div class='info-box'>
                                            <p><strong>File Name:</strong> {final_filename}</p>
                                            <p><strong>Uploaded By:</strong> {uploader_name}</p>
                                        </div>
                                        <p>You can access this file in the Document Management System.</p>
                                    </div>
                                    <div class='footer'>
                                        <p>&copy; 2025 DMS Portal. All rights reserved.</p>
                                    </div>
                                </div>
                            </body>
                            </html>
                            """
                            
                            message = MIMEMultipart("alternative")
                            message["Subject"] = subject
                            message["From"] = f"DOCUMENT MANAGEMENT SYSTEM <{sender_email}>"
                            message["To"] = recipient_email
                            
                            part = MIMEText(html_content, "html")
                            message.attach(part)
                            
                            server = smtplib.SMTP(smtp_server, smtp_port)
                            server.starttls()
                            server.login(sender_email, sender_password)
                            server.sendmail(sender_email, recipient_email, message.as_string())
                            server.quit()
                            
                            print(f"Email sent successfully to {recipient_email}")
                        except Exception as email_error:
                            print(f"Error sending email to {dept_user.get('email')}: {str(email_error)}")
            except Exception as e:
                print(f"Error in email notification: {str(e)}")
                # Don't fail the upload if email fails

        # Automatically grant department-level access when a department is specified
        if department_id:
            try:
                dept_id_int = int(department_id)
                share_exists = DatabaseConfig.execute_query(
                    "SELECT share_id FROM file_shares WHERE file_id=%s AND shared_with_department_id=%s",
                    (file_id, dept_id_int),
                    fetch_one=True
                )
                if not share_exists:
                    share_query = """INSERT INTO file_shares 
                                     (file_id, shared_with_user_id, shared_with_department_id, permission, shared_by)
                                     VALUES (%s, %s, %s, %s, %s)"""
                    DatabaseConfig.execute_query(
                        share_query,
                        (file_id, None, dept_id_int, 'viewer', user_id)
                    )
                    DMSDatabase.update_file(file_id, {'shared': True})
            except Exception as share_error:
                print(f"Department share creation error: {share_error}")

        if file_id:
            print(f"DEBUG: File created successfully with ID: {file_id}")
            
            # Add file to selected workspaces (many-to-many)
            workspace_additions = []
            for workspace_id in workspace_ids:
                if workspace_id and workspace_id != 'null':
                    workspace_id_int = int(workspace_id)
                    success = DMSDatabase.add_file_to_workspace(file_id, workspace_id_int)
                    workspace_additions.append({
                        'workspace_id': workspace_id_int,
                        'success': success
                    })
                    print(f"DEBUG: Added file {file_id} to workspace {workspace_id_int}: {success}")
            
            return jsonify({
                'success': True,
                'file_id': file_id,
                'filename': final_filename,
                'original_name': file.filename,
                'workspace_additions': workspace_additions,
                'file_path': file_path_hash,
                'classification': classification_result
            })
        else:
            # Clean up the saved file if database insertion failed
            print(f"DEBUG: Database insertion failed, cleaning up file")
            if os.path.exists(file_path):
                os.remove(file_path)
            return jsonify({'error': 'Failed to save file information to database'}), 500
        
    except Exception as e:
        print(f"DEBUG: Upload failed with exception: {str(e)}")
        return jsonify({'error': f'Upload failed: {str(e)}'}), 500
    
    
# Fix the download route in app.py
@app.route('/api/download-file/<int:file_id>', methods=['GET'])
def download_file(file_id):
    try:
        file_record = DMSDatabase.get_file_by_id(file_id)
        if not file_record:
            return jsonify({'error': 'File not found in database'}), 404
        
        request_user = get_request_user()
        if not request_user:
            return jsonify({'error': 'Unauthorized'}), 401
        
        permission = get_user_file_permission(file_record, request_user)
        if permission not in ('owner', 'editor', 'admin', 'viewer'):
            return jsonify({'error': 'Access denied'}), 403
        
        file_path = resolve_file_path(file_record)
        print(f"DEBUG: Download requested for file_id {file_id}")
        print(f"DEBUG: Resolved file path: {file_path}")
        
        if not file_path or not os.path.exists(file_path):
            return jsonify({'error': 'File not found on server. Please re-upload the file.'}), 404
        
        # Use the current stored `name` for download (this is the visible/current filename)
        # Fallback to `original_name` if `name` is missing
        download_name = file_record.get('name') or file_record.get('original_name')
        
        print(f"DEBUG: Sending file: {file_path} as {download_name}")
        return send_file(
            file_path,
            as_attachment=True,
            download_name=download_name
        )
        
    except Exception as e:
        print(f"Download error: {e}")
        return jsonify({'error': f'Download failed: {str(e)}'}), 500
    

# New route for file preview
# Visual file preview route - converts everything to PDF for exact visual representation
@app.route('/api/file-preview/<int:file_id>', methods=['GET'])
def file_preview(file_id):
    try:
        file_record = DMSDatabase.get_file_by_id(file_id)
        if not file_record:
            return jsonify({'error': 'File not found'}), 404
        
        request_user = get_request_user()
        if not request_user:
            return jsonify({'error': 'Unauthorized'}), 401
        
        permission = get_user_file_permission(file_record, request_user)
        if not permission:
            return jsonify({'error': 'Access denied'}), 403
        
        actual_file_path = resolve_file_path(file_record)
        if not actual_file_path or not os.path.exists(actual_file_path):
            return jsonify({'error': 'File not found on server'}), 404

        file_type = file_record['type'].lower()
        file_name = file_record.get('original_name') or file_record.get('name')
        
        print(f"DEBUG: Visual preview requested for {file_name} (type: {file_type})")
        print(f"DEBUG: Actual file path resolved: {actual_file_path}")

        # For PDFs - serve directly inline (perfect visual)
        if file_type == 'pdf':
            return send_file(actual_file_path, mimetype='application/pdf', as_attachment=False)
        
        # For images - convert to PDF for consistent preview
        elif file_type in ['png', 'jpg', 'jpeg', 'gif', 'bmp']:
            try:
                pdf_path = convert_image_to_pdf(actual_file_path)
                if pdf_path and os.path.exists(pdf_path):
                    return send_file(pdf_path, mimetype='application/pdf', as_attachment=False)
                else:
                    # Fallback: serve image directly inline
                    # Normalize common image mimetypes
                    image_mimetypes = {
                        'jpg': 'image/jpeg',
                        'jpeg': 'image/jpeg',
                        'png': 'image/png',
                        'gif': 'image/gif',
                        'bmp': 'image/bmp'
                    }
                    mimetype = image_mimetypes.get(file_type, f'image/{file_type}')
                    return send_file(actual_file_path, mimetype=mimetype, as_attachment=False)
            except Exception as e:
                print(f"Image to PDF conversion error: {e}")
                image_mimetypes = {
                    'jpg': 'image/jpeg',
                    'jpeg': 'image/jpeg',
                    'png': 'image/png',
                    'gif': 'image/gif',
                    'bmp': 'image/bmp'
                }
                mimetype = image_mimetypes.get(file_type, f'image/{file_type}')
                return send_file(actual_file_path, mimetype=mimetype, as_attachment=False)
        
        # For DOCX documents - convert to PDF for visual parity
        elif file_type == 'docx':
            try:
                pdf_path = convert_word_to_pdf_visual(actual_file_path, file_type)
                if pdf_path and os.path.exists(pdf_path):
                    return send_file(pdf_path, mimetype='application/pdf', as_attachment=False)
                else:
                    return create_fallback_preview(actual_file_path, file_type, "Word Document")
            except Exception as e:
                print(f"DOCX preview error: {e}")
                return create_fallback_preview(actual_file_path, file_type, "Word Document")
        
        # For legacy DOC files - fall back to PDF conversion
        elif file_type == 'doc':
            try:
                pdf_path = convert_word_to_pdf_visual(actual_file_path, file_type)
                if pdf_path and os.path.exists(pdf_path):
                    return send_file(pdf_path, mimetype='application/pdf', as_attachment=False)
                else:
                    return create_fallback_preview(actual_file_path, file_type, "Word Document")
            except Exception as e:
                print(f"Word to PDF conversion error: {e}")
                return create_fallback_preview(actual_file_path, file_type, "Word Document")
        
        # For Excel files - convert to PDF for real preview
        elif file_type in ['xls', 'xlsx']:
            try:
                pdf_path = convert_excel_to_pdf_visual(actual_file_path, file_type)
                if pdf_path and os.path.exists(pdf_path):
                    return send_file(pdf_path, mimetype='application/pdf', as_attachment=False)
                else:
                    return create_fallback_preview(actual_file_path, file_type, "Excel Spreadsheet")
            except Exception as e:
                print(f"Excel preview error: {e}")
                return create_fallback_preview(actual_file_path, file_type, "Excel Spreadsheet")
        
        # For PowerPoint files - convert to PDF for exact visual
        elif file_type in ['ppt', 'pptx']:
            try:
                pdf_path = convert_powerpoint_to_pdf_visual(actual_file_path, file_type)
                if pdf_path and os.path.exists(pdf_path):
                    return send_file(pdf_path, mimetype='application/pdf', as_attachment=False)
                else:
                    return create_fallback_preview(actual_file_path, file_type, "PowerPoint Presentation")
            except Exception as e:
                print(f"PowerPoint to PDF conversion error: {e}")
                return create_fallback_preview(actual_file_path, file_type, "PowerPoint Presentation")
        
        # For text files - create formatted PDF
        elif file_type == 'txt':
            try:
                pdf_path = convert_text_to_pdf(actual_file_path)
                if pdf_path and os.path.exists(pdf_path):
                    return send_file(pdf_path, mimetype='application/pdf', as_attachment=False)
                else:
                    return create_fallback_preview(actual_file_path, file_type, "Text File")
            except Exception as e:
                print(f"Text to PDF conversion error: {e}")
                return create_fallback_preview(actual_file_path, file_type, "Text File")
        
        # For ZIP files - create informative PDF
        elif file_type == 'zip':
            try:
                pdf_path = convert_zip_to_pdf(actual_file_path)
                if pdf_path and os.path.exists(pdf_path):
                    return send_file(pdf_path, mimetype='application/pdf', as_attachment=False)
                else:
                    return create_fallback_preview(actual_file_path, file_type, "ZIP Archive")
            except Exception as e:
                print(f"ZIP to PDF conversion error: {e}")
                return create_fallback_preview(actual_file_path, file_type, "ZIP Archive")
        
        # Unsupported file types
        else:
            return create_unsupported_preview(file_type)
            
    except Exception as e:
        return jsonify({'error': f'Preview failed: {str(e)}'}), 500
    
import tempfile
import os
from PIL import Image
import zipfile

def convert_image_to_pdf(image_path):
    """Convert image to PDF for perfect visual preview"""
    try:
        pdf_path = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf').name
        
        # Open image and convert to PDF
        with Image.open(image_path) as img:
            # Convert to RGB if necessary (for JPEG)
            if img.mode in ('RGBA', 'P'):
                img = img.convert('RGB')
            
            # Save as PDF
            img.save(pdf_path, 'PDF', resolution=100.0)
        
        print(f"DEBUG: Image converted to PDF: {pdf_path}")
        return pdf_path
        
    except Exception as e:
        print(f"Image to PDF conversion error: {e}")
        return None

def convert_word_to_pdf_visual(file_path, file_type):
    """Convert Word documents to PDF using LibreOffice (or fallback to COM)"""
    try:
        pdf_path = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf').name
        # First try LibreOffice (most reliable)
        try:
            import subprocess, shutil
            output_dir = os.path.dirname(pdf_path)
            os.makedirs(output_dir, exist_ok=True)

            # Try common LibreOffice locations and PATH
            soffice_candidates = [
                r"C:\Program Files\LibreOffice\program\soffice.exe",
                r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
            ]
            soffice_path = None
            for candidate in soffice_candidates:
                if os.path.exists(candidate):
                    soffice_path = candidate
                    break

            if not soffice_path:
                soffice_path = shutil.which('soffice')

            if soffice_path:
                cmd = [soffice_path, "--headless", "--convert-to", "pdf", "--outdir", output_dir, file_path]
                result = subprocess.run(cmd, capture_output=True, timeout=60)

                base_name = os.path.splitext(os.path.basename(file_path))[0]
                libreoffice_pdf = os.path.join(output_dir, base_name + ".pdf")

                if result.returncode == 0 and os.path.exists(libreoffice_pdf):
                    os.rename(libreoffice_pdf, pdf_path)
                    print(f"DEBUG: Word document converted to PDF via LibreOffice: {pdf_path}")
                    return pdf_path
                else:
                    print("DEBUG: LibreOffice returned non-zero or PDF not found")
                    print(f"DEBUG: soffice_path={soffice_path}, returncode={getattr(result, 'returncode', None)}")
                    print(f"DEBUG: stdout={getattr(result, 'stdout', b'').decode('utf-8', 'ignore')}")
                    print(f"DEBUG: stderr={getattr(result, 'stderr', b'').decode('utf-8', 'ignore')}")
            else:
                print("DEBUG: LibreOffice 'soffice' binary not found on system PATH or common locations")
        except Exception as lo_error:
            print(f"LibreOffice conversion failed: {lo_error}")
        
        # Fallback to Windows COM if LibreOffice not available
        try:
            import comtypes.client
            word = comtypes.client.CreateObject("Word.Application")
            word.Visible = False

            # Open document (read-only recommended)
            doc = word.Documents.Open(os.path.abspath(file_path), ReadOnly=1)
            doc.SaveAs(os.path.abspath(pdf_path), FileFormat=17)  # 17 = wdFormatPDF
            doc.Close()
            word.Quit()
            print(f"DEBUG: Word document converted to PDF via COM: {pdf_path}")
            return pdf_path
        except Exception as com_error:
            print(f"COM conversion also failed: {com_error}")
            return None
            
    except Exception as e:
        print(f"Word to PDF conversion error: {e}")
        return None

def convert_excel_to_pdf_visual(file_path, file_type):
    """Convert Excel files to PDF using LibreOffice (or fallback to COM)"""
    try:
        pdf_path = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf').name
        # First try LibreOffice
        try:
            import subprocess, shutil
            output_dir = os.path.dirname(pdf_path)
            os.makedirs(output_dir, exist_ok=True)

            soffice_candidates = [
                r"C:\Program Files\LibreOffice\program\soffice.exe",
                r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
            ]
            soffice_path = None
            for candidate in soffice_candidates:
                if os.path.exists(candidate):
                    soffice_path = candidate
                    break
            if not soffice_path:
                soffice_path = shutil.which('soffice')

            if soffice_path:
                cmd = [soffice_path, "--headless", "--convert-to", "pdf", "--outdir", output_dir, file_path]
                result = subprocess.run(cmd, capture_output=True, timeout=60)

                base_name = os.path.splitext(os.path.basename(file_path))[0]
                libreoffice_pdf = os.path.join(output_dir, base_name + ".pdf")

                if result.returncode == 0 and os.path.exists(libreoffice_pdf):
                    os.rename(libreoffice_pdf, pdf_path)
                    print(f"DEBUG: Excel file converted to PDF via LibreOffice: {pdf_path}")
                    return pdf_path
                else:
                    print("DEBUG: LibreOffice returned non-zero or PDF not found for Excel")
                    print(f"DEBUG: stdout={getattr(result, 'stdout', b'').decode('utf-8', 'ignore')}")
                    print(f"DEBUG: stderr={getattr(result, 'stderr', b'').decode('utf-8', 'ignore')}")
            else:
                print("DEBUG: LibreOffice 'soffice' binary not found on system PATH or common locations")
        except Exception as lo_error:
            print(f"LibreOffice conversion failed: {lo_error}")
        
        # Fallback to Windows COM
        try:
            import comtypes.client

            excel = comtypes.client.CreateObject("Excel.Application")
            excel.Visible = False

            workbook = excel.Workbooks.Open(os.path.abspath(file_path))
            workbook.ExportAsFixedFormat(0, os.path.abspath(pdf_path))  # 0 = xlTypePDF
            workbook.Close()
            excel.Quit()

            print(f"DEBUG: Excel workbook converted to PDF via COM: {pdf_path}")
            return pdf_path
        except Exception as com_error:
            print(f"COM conversion also failed: {com_error}")
            return None
            
    except Exception as e:
        print(f"Excel to PDF conversion error: {e}")
        return None

def convert_powerpoint_to_pdf_visual(file_path, file_type):
    """Convert PowerPoint to PDF using LibreOffice (or fallback to COM)"""
    try:
        pdf_path = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf').name
        # First try LibreOffice
        try:
            import subprocess, shutil
            output_dir = os.path.dirname(pdf_path)
            os.makedirs(output_dir, exist_ok=True)

            soffice_candidates = [
                r"C:\Program Files\LibreOffice\program\soffice.exe",
                r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
            ]
            soffice_path = None
            for candidate in soffice_candidates:
                if os.path.exists(candidate):
                    soffice_path = candidate
                    break
            if not soffice_path:
                soffice_path = shutil.which('soffice')

            if soffice_path:
                cmd = [soffice_path, "--headless", "--convert-to", "pdf", "--outdir", output_dir, file_path]
                result = subprocess.run(cmd, capture_output=True, timeout=60)

                base_name = os.path.splitext(os.path.basename(file_path))[0]
                libreoffice_pdf = os.path.join(output_dir, base_name + ".pdf")

                if result.returncode == 0 and os.path.exists(libreoffice_pdf):
                    os.rename(libreoffice_pdf, pdf_path)
                    print(f"DEBUG: PowerPoint file converted to PDF via LibreOffice: {pdf_path}")
                    return pdf_path
                else:
                    print("DEBUG: LibreOffice returned non-zero or PDF not found for PowerPoint")
                    print(f"DEBUG: stdout={getattr(result, 'stdout', b'').decode('utf-8', 'ignore')}")
                    print(f"DEBUG: stderr={getattr(result, 'stderr', b'').decode('utf-8', 'ignore')}")
            else:
                print("DEBUG: LibreOffice 'soffice' binary not found on system PATH or common locations")
        except Exception as lo_error:
            print(f"LibreOffice conversion failed: {lo_error}")
        
        # Fallback to Windows COM
        try:
            import comtypes.client

            powerpoint = comtypes.client.CreateObject("PowerPoint.Application")
            powerpoint.Visible = False

            presentation = powerpoint.Presentations.Open(os.path.abspath(file_path), WithWindow=False)
            presentation.SaveAs(os.path.abspath(pdf_path), 32)  # 32 = ppSaveAsPDF
            presentation.Close()
            powerpoint.Quit()

            print(f"DEBUG: PowerPoint file converted to PDF via COM: {pdf_path}")
            return pdf_path
        except Exception as com_error:
            print(f"COM conversion also failed: {com_error}")
            return None
            
    except Exception as e:
        print(f"PowerPoint to PDF conversion error: {e}")
        return None

def convert_text_to_pdf(file_path):
    """Convert text file to formatted PDF"""
    try:
        pdf_path = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf').name
        
        # Read text content
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            content = f.read()
        
        # Create PDF with formatted text
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.units import inch
        
        c = canvas.Canvas(pdf_path, pagesize=letter)
        c.setFont("Courier", 10)  # Monospace font for text files
        
        # Simple text wrapping
        lines = content.split('\n')
        y_position = 750
        line_height = 12
        
        for line in lines[:100]:  # Limit to first 100 lines
            if y_position < 50:
                c.showPage()
                c.setFont("Courier", 10)
                y_position = 750
            
            # Truncate long lines
            if len(line) > 120:
                line = line[:117] + "..."
            
            c.drawString(50, y_position, line)
            y_position -= line_height
        
        c.save()
        print(f"DEBUG: Text file converted to PDF: {pdf_path}")
        return pdf_path
        
    except Exception as e:
        print(f"Text to PDF conversion error: {e}")
        return None

def convert_zip_to_pdf(file_path):
    """Create PDF showing ZIP file contents"""
    try:
        pdf_path = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf').name
        
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import letter
        
        c = canvas.Canvas(pdf_path, pagesize=letter)
        c.setFont("Helvetica-Bold", 16)
        c.drawString(50, 750, "ZIP File Contents")
        
        c.setFont("Helvetica", 10)
        y_position = 720
        
        # List files in ZIP
        with zipfile.ZipFile(file_path, 'r') as zip_ref:
            file_list = zip_ref.namelist()
        
        for i, file_name in enumerate(file_list[:50]):  # First 50 files
            if y_position < 50:
                c.showPage()
                c.setFont("Helvetica", 10)
                y_position = 750
            
            c.drawString(70, y_position, f"{i+1}. {file_name}")
            y_position -= 15
        
        # Add summary
        if len(file_list) > 50:
            y_position -= 10
            c.drawString(70, y_position, f"... and {len(file_list) - 50} more files")
        
        c.save()
        print(f"DEBUG: ZIP file converted to PDF: {pdf_path}")
        return pdf_path
        
    except Exception as e:
        print(f"ZIP to PDF conversion error: {e}")
        return None

def docx_to_html(path):
    doc = DocxDocument(path)
    html_parts = []
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            html_parts.append(f"<p>{escape(text)}</p>")
    if doc.tables:
        for table in doc.tables:
            html_parts.append("<table class='preview-table'>")
            for row in table.rows:
                html_parts.append("<tr>")
                for cell in row.cells:
                    cell_text = escape(cell.text.strip())
                    html_parts.append(f"<td>{cell_text}</td>")
                html_parts.append("</tr>")
            html_parts.append("</table>")
    if not html_parts:
        html_parts.append("<p>(Document is empty)</p>")
    return "".join(html_parts)

def excel_to_html(path):
    sheets = pd.read_excel(path, sheet_name=None, dtype=str, engine=None)
    html_parts = []
    for sheet_name, df in sheets.items():
        df = df.fillna("")
        html_parts.append(f"<h4>{escape(str(sheet_name))}</h4>")
        html_parts.append(df.to_html(classes="preview-table", index=False, escape=True))
    if not html_parts:
        html_parts.append("<p>(Workbook is empty)</p>")
    return "".join(html_parts)

def create_simple_pdf_from_content(file_path, file_type):
    """Create a simple PDF when COM conversion fails"""
    try:
        pdf_path = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf').name
        
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import letter
        
        c = canvas.Canvas(pdf_path, pagesize=letter)
        
        # Header
        c.setFont("Helvetica-Bold", 16)
        c.drawString(50, 750, f"{file_type} Preview")
        c.setFont("Helvetica", 10)
        c.drawString(50, 730, f"File: {os.path.basename(file_path)}")
        
        # Content message
        c.drawString(50, 700, "This file is being displayed in preview mode.")
        c.drawString(50, 680, "For exact visual representation, please download the file.")
        
        # File info
        file_size = os.path.getsize(file_path)
        c.drawString(50, 650, f"File size: {file_size} bytes")
        
        c.save()
        return pdf_path
        
    except Exception as e:
        print(f"Simple PDF creation error: {e}")
        return None

def create_fallback_preview(file_path, file_type, description):
    """Create fallback PDF preview"""
    pdf_path = create_simple_pdf_from_content(file_path, description)
    if pdf_path and os.path.exists(pdf_path):
        return send_file(pdf_path, mimetype='application/pdf', as_attachment=False)
    else:
        return jsonify({
            'type': 'unsupported',
            'message': f'Preview not available for {file_type.upper()} files'
        })

def create_unsupported_preview(file_type):
    """Create PDF for unsupported file types"""
    try:
        pdf_path = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf').name
        
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import letter
        
        c = canvas.Canvas(pdf_path, pagesize=letter)
        c.setFont("Helvetica-Bold", 16)
        c.drawString(50, 750, "File Preview Not Available")
        c.setFont("Helvetica", 12)
        c.drawString(50, 720, f"File type: {file_type.upper()}")
        c.drawString(50, 700, "This file type cannot be previewed in the system.")
        c.drawString(50, 680, "Please download the file to view its contents.")
        
        c.save()
        return send_file(pdf_path, mimetype='application/pdf', as_attachment=False)
        
    except Exception as e:
        return jsonify({
            'type': 'unsupported',
            'message': f'Preview not available for {file_type.upper()} files'
        })

class DMSDatabase:
    # User operations
    @staticmethod
    def get_user_by_id(user_id):
        query = "SELECT * FROM users WHERE user_id = %s"
        return DatabaseConfig.execute_query(query, (user_id,), fetch_one=True)

    @staticmethod
    def get_user_by_email(email):
        query = "SELECT * FROM users WHERE email = %s"
        return DatabaseConfig.execute_query(query, (email,), fetch_one=True)

    @staticmethod
    def get_all_users():
        query = "SELECT u.*, d.name as department_name FROM users u LEFT JOIN departments d ON u.department_id = d.department_id"
        return DatabaseConfig.execute_query(query, fetch=True)

    @staticmethod
    def create_user(name, email, password, role, department_id):
        # Get the maximum existing user_id and increment by 1
        max_id_query = "SELECT MAX(user_id) as max_id FROM users"
        result = DatabaseConfig.execute_query(max_id_query, fetch_one=True)
        next_id = (result['max_id'] + 1) if result['max_id'] else 1
        
        query = """INSERT INTO users (user_id, name, email, password, role, department_id) 
                VALUES (%s, %s, %s, %s, %s, %s)"""
        return DatabaseConfig.execute_query(query, (next_id, name, email, password, role, department_id), lastrowid=True)

    @staticmethod
    def update_user(user_id, name=None, email=None, role=None, department_id=None, status=None, profile_image=None, password=None, clear_department=False):
        """Update user information - supports partial updates"""
        try:
            connection = DatabaseConfig.get_connection()
            cursor = connection.cursor()
            
            # Build update query dynamically based on what's provided
            updates = []
            params = []
            
            if name is not None:
                updates.append("name=%s")
                params.append(name)
            if email is not None:
                updates.append("email=%s")
                params.append(email)
            if role is not None:
                updates.append("role=%s")
                params.append(role)
            # Handle department changes. DB enforces NOT NULL on users.department_id,
            # so when clear_department is requested, map to an "Unassigned" department.
            if department_id is not None or clear_department:
                target_department_id = department_id
                if clear_department and (department_id is None):
                    # Find or create a catch-all department for unassigned users
                    try:
                        # Try to find existing 'Unassigned' department
                        find_sql = "SELECT department_id FROM departments WHERE name=%s"
                        cursor.execute(find_sql, ("Unassigned",))
                        row = cursor.fetchone()
                        if row and row[0]:
                            target_department_id = row[0]
                        else:
                            # Create 'Unassigned' department owned by the user being updated
                            # Compute next department_id similar to create_department
                            cursor.execute("SELECT COALESCE(MAX(department_id), 4000) + 1 FROM departments")
                            next_dept_id = cursor.fetchone()[0]
                            insert_sql = "INSERT INTO departments (department_id, name, user_id) VALUES (%s, %s, %s)"
                            cursor.execute(insert_sql, (next_dept_id, "Unassigned", user_id))
                            target_department_id = next_dept_id
                            print(f"DEBUG: Created fallback 'Unassigned' department with id {next_dept_id}")
                    except Exception as dept_err:
                        print(f"WARNING: Failed ensuring 'Unassigned' department: {dept_err}")
                        # As a last resort, try to fall back to any existing department
                        cursor.execute("SELECT department_id FROM departments ORDER BY department_id LIMIT 1")
                        any_row = cursor.fetchone()
                        if not any_row:
                            raise
                        target_department_id = any_row[0]

                updates.append("department_id=%s")
                params.append(target_department_id)
            if status is not None:
                updates.append("status=%s")
                params.append(status)
            if profile_image is not None:
                updates.append("profile_image=%s")
                params.append(profile_image)
            if password is not None:
                updates.append("password=%s")
                params.append(password)
            
            if not updates:
                cursor.close()
                connection.close()
                return False
            
            updates.append("updated_at=NOW()")
            params.append(user_id)
            
            query = f"UPDATE users SET {', '.join(updates)} WHERE user_id=%s"
            
            cursor.execute(query, params)
            connection.commit()
            affected_rows = cursor.rowcount
            cursor.close()
            connection.close()
            
            print(f"DEBUG: Updated user {user_id}, affected rows: {affected_rows}")
            # Treat successful execution as success even if rowcount is 0
            # (e.g., value already the same). Prevent false 500s on no-op updates.
            return True
            
        except Exception as e:
            print(f"Error updating user: {e}")
            if 'connection' in locals():
                connection.rollback()
            return False
        
    @staticmethod
    def archive_user(user_id, reason=None, notes=None):
        # Update status to archived (reason and notes are optional for backward compatibility)
        if reason or notes:
            query = """UPDATE users SET status='archived', archive_reason=%s, archive_notes=%s, archived_at=NOW() 
                       WHERE user_id=%s"""
            return DatabaseConfig.execute_query(query, (reason, notes, user_id))
        else:
            query = """UPDATE users SET status='archived', archived_at=NOW() 
                       WHERE user_id=%s"""
            return DatabaseConfig.execute_query(query, (user_id,))
    
    @staticmethod
    def delete_user(user_id):
        """Permanently delete a user and all related records"""
        try:
            connection = DatabaseConfig.get_connection()
            cursor = connection.cursor()
            
            try:
                # Delete in order to respect foreign key constraints
                cursor.execute("DELETE FROM file_shares WHERE shared_by = %s OR shared_with_user_id = %s", (user_id, user_id))
                cursor.execute("DELETE FROM workspace_members WHERE user_id = %s", (user_id,))
                cursor.execute("DELETE FROM sessions WHERE user_id = %s", (user_id,))
                cursor.execute("DELETE FROM activity_log WHERE user_id = %s", (user_id,))
                cursor.execute("DELETE FROM files WHERE user_id = %s", (user_id,))
                cursor.execute("DELETE FROM workspaces WHERE user_id = %s", (user_id,))
                cursor.execute("DELETE FROM users WHERE user_id = %s", (user_id,))
                
                connection.commit()
                cursor.close()
                connection.close()
                return True
            except Exception as e:
                connection.rollback()
                cursor.close()
                connection.close()
                print(f"Error deleting user {user_id}: {e}")
                import traceback
                traceback.print_exc()
                return False
        except Exception as e:
            print(f"Error deleting user: {e}")
            return False

    # File operations
    @staticmethod
    def get_all_files():
        query = """SELECT f.*, u.name as owner_name, c.name as category_name 
                   FROM files f 
                   LEFT JOIN users u ON f.user_id = u.user_id 
                   LEFT JOIN categories c ON f.category_id = c.category_id 
                   WHERE f.status = 'active'"""
        return DatabaseConfig.execute_query(query, fetch=True)

    @staticmethod
    def get_active_files_by_user(user_id):
        query = "SELECT * FROM files WHERE user_id = %s AND status = 'active'"
        return DatabaseConfig.execute_query(query, (user_id,), fetch=True)

    @staticmethod
    def get_file_by_id(file_id):
        query = "SELECT * FROM files WHERE file_id = %s"
        return DatabaseConfig.execute_query(query, (file_id,), fetch_one=True)

    @staticmethod
    def create_file(file_data):
        query = """INSERT INTO files (name, original_name, size, type, user_id, workspace_id, category_id, 
                document_type, classification_confidence, classification_error, text_sample, 
                shared, file_path, is_encrypted, encryption_version) 
                VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)"""
        params = (
            file_data['name'], file_data['original_name'], file_data['size'], file_data['type'],
            file_data['user_id'], file_data.get('workspace_id'), file_data.get('category_id'),
            file_data.get('document_type'), file_data.get('classification_confidence'),
            file_data.get('classification_error'), file_data.get('text_sample'),
            file_data.get('shared', False),
            file_data.get('file_path'), file_data.get('is_encrypted', False),
            file_data.get('encryption_version')
        )
        return DatabaseConfig.execute_query(query, params, lastrowid=True)

    
    @staticmethod
    def update_file(file_id, updates):
        try:
            set_clause = ", ".join([f"{key}=%s" for key in updates.keys()])
            query = f"UPDATE files SET {set_clause} WHERE file_id=%s"
            params = list(updates.values()) + [file_id]
            return DatabaseConfig.execute_query(query, params)
        except Exception as e:
            print(f"Error updating file: {e}")
            return None

    @staticmethod
    def delete_file(file_id):
        connection = None
        cursor = None
        try:
            file_record = DMSDatabase.get_file_by_id(file_id)
            if not file_record:
                print(f"DEBUG: File {file_id} not found for deletion")
                return False
            
            print(f"DEBUG: Deleting file {file_id} - {file_record.get('name')}")
            
            if file_record.get('file_path'):
                try:
                    actual_path = resolve_file_path(file_record)
                    if actual_path and os.path.exists(actual_path):
                        os.remove(actual_path)
                        print(f"DEBUG: Physical file deleted: {actual_path}")
                except Exception as e:
                    print(f"Warning: Could not delete physical file: {e}")
            
            connection = DatabaseConfig.get_connection()
            cursor = connection.cursor()
            
            cursor.execute("DELETE FROM file_workspaces WHERE file_id = %s", (file_id,))
            cursor.execute("DELETE FROM file_shares WHERE file_id = %s", (file_id,))
            cursor.execute("DELETE FROM files WHERE file_id = %s", (file_id,))
            
            connection.commit()
            print(f"DEBUG: File {file_id} database deletion completed successfully")
            return True
            
        except Exception as e:
            print(f"Error deleting file {file_id}: {e}")
            if connection:
                connection.rollback()
            return False
        finally:
            if cursor:
                cursor.close()
            if connection:
                connection.close()

    @staticmethod
    def archive_file(file_id):
        try:
            print(f"DEBUG: Starting archive process for file {file_id}")
            
            # First get the file details
            file_record = DMSDatabase.get_file_by_id(file_id)
            if not file_record:
                print(f"DEBUG: File {file_id} not found in database")
                return False
            
            print(f"DEBUG: Found file: {file_record['name']} (current status: {file_record.get('status', 'unknown')})")
            
            # If file is already archived, return True
            if file_record.get('status') == 'archived':
                print(f"DEBUG: File {file_id} is already archived")
                return True
            
            # Generate unique name for archive
            original_name = file_record['original_name']
            user_id = file_record['user_id']
            
            # Get unique filename for archive
            base_name = original_name.rsplit('.', 1)[0]
            extension = original_name.rsplit('.', 1)[1].lower()
            
            query = "SELECT name FROM files WHERE user_id = %s AND status = 'archived' AND name LIKE %s"
            connection = DatabaseConfig.get_connection()
            cursor = connection.cursor(dictionary=True)
            
            like_pattern = f"{base_name}%.{extension}"
            cursor.execute(query, (user_id, like_pattern))
            existing_files = cursor.fetchall()
            cursor.close()
            connection.close()
            
            final_filename = original_name
            counter = 1
            
            existing_names = [f['name'] for f in existing_files]
            while final_filename in existing_names:
                final_filename = f"{base_name} ({counter}).{extension}"
                counter += 1
            
            print(f"DEBUG: Updating file {file_id} to archived with name: {final_filename}")
            
            # Update status and name for archive - USE DIRECT DATABASE CONNECTION
            query = "UPDATE files SET status='archived', name=%s WHERE file_id=%s"
            connection = DatabaseConfig.get_connection()
            cursor = connection.cursor()
            cursor.execute(query, (final_filename, file_id))
            affected_rows = cursor.rowcount
            connection.commit()
            cursor.close()
            connection.close()
            
            print(f"DEBUG: Archive query affected {affected_rows} rows")
            
            # Verify the update worked by checking the file status
            updated_file = DMSDatabase.get_file_by_id(file_id)
            if updated_file and updated_file.get('status') == 'archived':
                print(f"DEBUG: File {file_id} successfully archived - status confirmed")
                return True
            else:
                print(f"DEBUG: File {file_id} archive failed - status not updated")
                return False
                
        except Exception as e:
            print(f"Error in archive_file for file {file_id}: {str(e)}")
            return False
        
    @staticmethod
    def is_file_in_workspace(file_id, workspace_id):
        """Check if file is already in a specific workspace"""
        try:
            query = "SELECT file_id FROM files WHERE file_id = %s AND workspace_id = %s AND status = 'active'"
            result = DatabaseConfig.execute_query(query, (file_id, workspace_id), fetch_one=True)
            return result is not None
        except Exception as e:
            print(f"Error checking file in workspace: {e}")
            return False

# ========== WORKSPACE METHODS - COMPLETE REPLACEMENT ==========

    @staticmethod
    def add_file_to_workspace(file_id, workspace_id):
        """Add file to a workspace (many-to-many relationship)"""
        try:
            print(f"🚀 DEBUG: Adding file {file_id} to workspace {workspace_id}")
            
            # Check if relationship already exists
            check_query = "SELECT * FROM file_workspaces WHERE file_id = %s AND workspace_id = %s"
            existing = DatabaseConfig.execute_query(check_query, (file_id, workspace_id), fetch_one=True)
            
            if existing:
                print(f"✅ DEBUG: File {file_id} already in workspace {workspace_id}")
                return True
                
            # Add new relationship
            query = "INSERT INTO file_workspaces (file_id, workspace_id) VALUES (%s, %s)"
            result = DatabaseConfig.execute_query(query, (file_id, workspace_id))
            print(f"✅ DEBUG: Added file {file_id} to workspace {workspace_id}: {result is not None}")
            return result is not None
        except Exception as e:
            print(f"❌ Error adding file to workspace: {e}")
            return False

    @staticmethod
    def remove_file_from_workspace(file_id, workspace_id):
        """Remove file from a specific workspace"""
        try:
            print(f"🚀 DEBUG: Removing file {file_id} from workspace {workspace_id}")
            query = "DELETE FROM file_workspaces WHERE file_id = %s AND workspace_id = %s"
            result = DatabaseConfig.execute_query(query, (file_id, workspace_id))
            print(f"✅ DEBUG: Removed file {file_id} from workspace {workspace_id}: {result is not None}")
            return result is not None
        except Exception as e:
            print(f"❌ Error removing file from workspace: {e}")
            return False

    @staticmethod
    def get_file_workspaces(file_id):
        """Get all workspaces where a file exists"""
        try:
            query = """SELECT w.workspace_id, w.name, w.description 
                    FROM workspaces w 
                    JOIN file_workspaces fw ON w.workspace_id = fw.workspace_id 
                    WHERE fw.file_id = %s"""
            workspaces = DatabaseConfig.execute_query(query, (file_id,), fetch=True) or []
            print(f"📁 DEBUG: File {file_id} is in {len(workspaces)} workspaces: {[w['workspace_id'] for w in workspaces]}")
            return workspaces
        except Exception as e:
            print(f"❌ Error getting file workspaces: {e}")
            return []

    @staticmethod
    def get_files_by_workspace(workspace_id):
        """Get all files in a specific workspace - USING NEW MANY-TO-MANY"""
        try:
            print(f"🚀 DEBUG: Getting files for workspace {workspace_id} from file_workspaces table")
            
            query = """SELECT f.*, u.name as owner_name 
                    FROM files f 
                    JOIN file_workspaces fw ON f.file_id = fw.file_id 
                    LEFT JOIN users u ON f.user_id = u.user_id 
                    WHERE fw.workspace_id = %s AND f.status = 'active'"""
            files = DatabaseConfig.execute_query(query, (workspace_id,), fetch=True) or []
            
            print(f"✅ DEBUG: Found {len(files)} files in workspace {workspace_id}")
            for file in files:
                print(f"   - File: {file['file_id']} - {file['name']}")
            
            return files
        except Exception as e:
            print(f"❌ Error getting files by workspace: {e}")
            return []

    @staticmethod
    def update_file_workspaces(file_id, workspace_ids):
        """Update all workspace relationships for a file"""
        try:
            print(f"🚀 DEBUG: Updating workspaces for file {file_id} to: {workspace_ids}")
            
            connection = DatabaseConfig.get_connection()
            cursor = connection.cursor()
            
            # Get current workspaces
            current_query = "SELECT workspace_id FROM file_workspaces WHERE file_id = %s"
            cursor.execute(current_query, (file_id,))
            current_workspaces = [row[0] for row in cursor.fetchall()]
            
            print(f"📊 DEBUG: Current workspaces: {current_workspaces}")
            print(f"📊 DEBUG: New workspaces: {workspace_ids}")
            
            # Workspaces to add
            to_add = [wid for wid in workspace_ids if wid not in current_workspaces]
            # Workspaces to remove  
            to_remove = [wid for wid in current_workspaces if wid not in workspace_ids]
            
            print(f"➕ DEBUG: To add: {to_add}")
            print(f"➖ DEBUG: To remove: {to_remove}")
            
            # Add new relationships
            added_count = 0
            for workspace_id in to_add:
                insert_query = "INSERT IGNORE INTO file_workspaces (file_id, workspace_id) VALUES (%s, %s)"
                cursor.execute(insert_query, (file_id, workspace_id))
                added_count += 1
                print(f"✅ DEBUG: Added file {file_id} to workspace {workspace_id}")
            
            # Remove old relationships
            removed_count = 0
            for workspace_id in to_remove:
                delete_query = "DELETE FROM file_workspaces WHERE file_id = %s AND workspace_id = %s"
                cursor.execute(delete_query, (file_id, workspace_id))
                removed_count += 1
                print(f"✅ DEBUG: Removed file {file_id} from workspace {workspace_id}")
            
            connection.commit()
            cursor.close()
            connection.close()
            
            print(f"🎉 DEBUG: Successfully updated - Added: {added_count}, Removed: {removed_count}")
            return True
        except Exception as e:
            print(f"❌ Error updating file workspaces: {e}")
            if 'connection' in locals():
                connection.rollback()
            return False

    @staticmethod
    def get_all_workspaces():
        """Get all workspaces with their files count"""
        try:
            query = """SELECT w.*, u.name as creator_name,
                            (SELECT COUNT(*) FROM file_workspaces fw WHERE fw.workspace_id = w.workspace_id) as file_count
                    FROM workspaces w 
                    LEFT JOIN users u ON w.user_id = u.user_id"""
            workspaces = DatabaseConfig.execute_query(query, fetch=True)
            
            print(f"🚀 DEBUG: Found {len(workspaces) if workspaces else 0} workspaces")
            
            # Add members to each workspace
            if workspaces:
                for workspace in workspaces:
                    members = DMSDatabase.get_workspace_members(workspace['workspace_id'])
                    workspace['members'] = [member['user_id'] for member in members] if members else []
                    print(f"👥 DEBUG: Workspace {workspace['workspace_id']} '{workspace['name']}' has {len(workspace['members'])} members and {workspace['file_count']} files")
            
            return workspaces
        except Exception as e:
            print(f"❌ Error getting workspaces: {e}")
            return []
        
    # Department operations
    @staticmethod
    def get_all_departments():
        query = "SELECT * FROM departments"
        return DatabaseConfig.execute_query(query, fetch=True)

    @staticmethod
    def create_department(name, user_id):
        # Get the maximum existing department_id and increment by 1
        max_id_query = "SELECT MAX(department_id) as max_id FROM departments"
        result = DatabaseConfig.execute_query(max_id_query, fetch_one=True)
        next_id = (result['max_id'] + 1) if result['max_id'] else 4001
        
        # Insert with explicit department_id
        query = "INSERT INTO departments (department_id, name, user_id) VALUES (%s, %s, %s)"
        DatabaseConfig.execute_query(query, (next_id, name, user_id))
        return next_id

    @staticmethod
    def delete_department(department_id):
        query = "DELETE FROM departments WHERE department_id = %s"
        return DatabaseConfig.execute_query(query, (department_id,))

    # Category operations
    @staticmethod
    def get_all_categories():
        query = "SELECT * FROM categories ORDER BY name"
        return DatabaseConfig.execute_query(query, fetch=True)

    @staticmethod
    def create_category(name, description=None, auto_created=False, is_unclassified=False):
        query = "INSERT INTO categories (name, description, auto_created, is_unclassified) VALUES (%s, %s, %s, %s)"
        return DatabaseConfig.execute_query(query, (name, description, auto_created, is_unclassified), lastrowid=True)

    @staticmethod
    def update_category(category_id, name, description):
        query = "UPDATE categories SET name=%s, description=%s WHERE category_id=%s"
        return DatabaseConfig.execute_query(query, (name, description, category_id))

    @staticmethod
    def get_category_by_name(name):
        query = "SELECT * FROM categories WHERE name = %s"
        return DatabaseConfig.execute_query(query, (name,), fetch_one=True)

    # Workspace operations
    @staticmethod
    def get_all_workspaces():
        try:
            query = """SELECT w.*, u.name as creator_name 
                    FROM workspaces w 
                    LEFT JOIN users u ON w.user_id = u.user_id"""
            workspaces = DatabaseConfig.execute_query(query, fetch=True)
            
            print(f"DEBUG: Found {len(workspaces) if workspaces else 0} workspaces")
            
            # Add members to each workspace
            if workspaces:
                for workspace in workspaces:
                    members = DMSDatabase.get_workspace_members(workspace['workspace_id'])
                    workspace['members'] = [member['user_id'] for member in members] if members else []
                    print(f"DEBUG: Workspace {workspace['workspace_id']} '{workspace['name']}' has members: {workspace['members']}")
            
            return workspaces
        except Exception as e:
            print(f"Error getting workspaces: {e}")
            return []

    @staticmethod
    def create_workspace(name, description, user_id, members):
        try:
            connection = DatabaseConfig.get_connection()
            cursor = connection.cursor(dictionary=True)
            
            # Get the maximum existing workspace_id
            cursor.execute("SELECT MAX(workspace_id) as max_id FROM workspaces")
            result = cursor.fetchone()
            next_id = (result['max_id'] + 1) if result and result['max_id'] else 2001
            
            print(f"DEBUG: Creating workspace with ID {next_id}, members: {members}")
            
            # Insert workspace
            query = "INSERT INTO workspaces (workspace_id, name, description, user_id) VALUES (%s, %s, %s, %s)"
            cursor.execute(query, (next_id, name, description, user_id))
            
            # Add members (ensure user_id is included and all are integers)
            all_members = set()
            if members:
                for member_id in members:
                    all_members.add(int(member_id))
            all_members.add(int(user_id))  # Always include creator
            
            print(f"DEBUG: Final members to add: {all_members}")
            
            for member_id in all_members:
                add_query = "INSERT IGNORE INTO workspace_members (workspace_id, user_id) VALUES (%s, %s)"
                cursor.execute(add_query, (next_id, member_id))
                print(f"DEBUG: Added member {member_id} to workspace {next_id}")
            
            connection.commit()
            cursor.close()
            connection.close()
            
            return next_id
        except Exception as e:
            print(f"Error creating workspace: {e}")
            if 'connection' in locals():
                connection.rollback()
            return None

    @staticmethod
    def add_workspace_member(workspace_id, user_id):
        query = "INSERT IGNORE INTO workspace_members (workspace_id, user_id) VALUES (%s, %s)"
        return DatabaseConfig.execute_query(query, (workspace_id, user_id))

    @staticmethod
    def get_workspace_members(workspace_id):
        try:
            query = """SELECT u.user_id, u.name, u.email 
                    FROM workspace_members wm 
                    JOIN users u ON wm.user_id = u.user_id 
                    WHERE wm.workspace_id = %s"""
            members = DatabaseConfig.execute_query(query, (workspace_id,), fetch=True)
            print(f"DEBUG: get_workspace_members for workspace {workspace_id} returned: {members}")
            return members
        except Exception as e:
            print(f"Error getting workspace members: {e}")
            return []
    
    # Add these methods to the DMSDatabase class

    @staticmethod
    def get_workspace_by_id(workspace_id):
        query = """SELECT w.*, u.name as creator_name 
                FROM workspaces w 
                LEFT JOIN users u ON w.user_id = u.user_id 
                WHERE w.workspace_id = %s"""
        return DatabaseConfig.execute_query(query, (workspace_id,), fetch_one=True)

    @staticmethod
    def update_workspace(workspace_id, name, description, members):
        try:
            connection = DatabaseConfig.get_connection()
            cursor = connection.cursor()
            
            # Update workspace basic info
            query = "UPDATE workspaces SET name=%s, description=%s, updated_at=NOW() WHERE workspace_id=%s"
            cursor.execute(query, (name, description, workspace_id))
            
            # Update members if provided
            if members is not None:
                print(f"DEBUG: Updating members for workspace {workspace_id}: {members}")
                
                # Clear current members
                delete_query = "DELETE FROM workspace_members WHERE workspace_id = %s"
                cursor.execute(delete_query, (workspace_id,))
                
                # Add new members (convert all to integers)
                for member_id in members:
                    member_id_int = int(member_id)
                    insert_query = "INSERT IGNORE INTO workspace_members (workspace_id, user_id) VALUES (%s, %s)"
                    cursor.execute(insert_query, (workspace_id, member_id_int))
                    print(f"DEBUG: Added member {member_id_int} to workspace {workspace_id}")
            
            connection.commit()
            cursor.close()
            connection.close()
            return True
        except Exception as e:
            print(f"Error updating workspace: {e}")
            if 'connection' in locals():
                connection.rollback()
            return False

    @staticmethod
    def delete_workspace(workspace_id):
        # First remove workspace members
        delete_members_query = "DELETE FROM workspace_members WHERE workspace_id = %s"
        DatabaseConfig.execute_query(delete_members_query, (workspace_id,))
        
        # Then delete the workspace
        query = "DELETE FROM workspaces WHERE workspace_id = %s"
        return DatabaseConfig.execute_query(query, (workspace_id,))

    @staticmethod
    def create_reset_code(email, code, token=None, expires_at=None):
        """
        Create a reset code for password reset.
        Stores email, code, token, expiry time, and tracks if it's been used.
        """
        query = "INSERT INTO reset_codes (email, code, token, expires_at, is_used) VALUES (%s, %s, %s, %s, %s)"
        import uuid
        import datetime
        if token is None:
            token = str(uuid.uuid4())
        if expires_at is None:
            # Token expires in 1 hour
            expires_at = datetime.datetime.now() + datetime.timedelta(hours=1)
        return DatabaseConfig.execute_query(query, (email, code, token, expires_at, 0), lastrowid=True)
    
    @staticmethod
    def verify_reset_code(email, code):
        """
        Verify if a reset code is valid (not used and not expired).
        """
        query = """SELECT * FROM reset_codes 
                   WHERE email = %s AND code = %s 
                   AND is_used = 0 
                   AND (expires_at IS NULL OR expires_at > NOW())
                   ORDER BY reset_codes_id DESC LIMIT 1"""
        return DatabaseConfig.execute_query(query, (email, code), fetch_one=True)
    
    @staticmethod
    def mark_reset_code_as_used(email, code):
        """
        Mark a reset code as used after successful password reset.
        """
        query = "UPDATE reset_codes SET is_used = 1 WHERE email = %s AND code = %s"
        return DatabaseConfig.execute_query(query, (email, code))

    # Settings operations
    @staticmethod
    def get_settings():
        query = "SELECT * FROM settings WHERE settings_id = 6001"
        settings = DatabaseConfig.execute_query(query, fetch_one=True)
        # Provide sane defaults if the row is missing or incomplete
        if not settings:
            return {
                'company': 'DMS',
                'max_file_mb': 50,
                'allowed_types': 'pdf,doc,docx,jpg,jpeg,png,txt',
                'company_logo': None,
            }
        # Ensure keys exist for frontend expectations
        settings.setdefault('company', 'DMS')
        settings.setdefault('max_file_mb', 50)
        settings.setdefault('allowed_types', 'pdf,doc,docx,jpg,jpeg,png,txt')
        settings.setdefault('company_logo', None)
        return settings

    @staticmethod
    def update_settings(company, max_file_mb, allowed_types, company_logo=None):
        """
        Update core system settings. If company_logo is None, the existing logo
        is preserved to avoid accidentally clearing it.
        """
        if company_logo is None:
            query = "UPDATE settings SET company=%s, max_file_mb=%s, allowed_types=%s WHERE settings_id=6001"
            params = (company, max_file_mb, allowed_types)
        else:
            query = "UPDATE settings SET company=%s, max_file_mb=%s, allowed_types=%s, company_logo=%s WHERE settings_id=6001"
            params = (company, max_file_mb, allowed_types, company_logo)
        return DatabaseConfig.execute_query(query, params)

    @staticmethod
    def update_company_logo(company_logo):
        query = "UPDATE settings SET company_logo=%s WHERE settings_id=6001"
        return DatabaseConfig.execute_query(query, (company_logo,))

    # Activity log operations
    @staticmethod
    def log_activity(user_id, activity):
        query = "INSERT INTO activity_log (user_id, activity) VALUES (%s, %s)"
        return DatabaseConfig.execute_query(query, (user_id, activity))

    @staticmethod
    def get_recent_activities(limit=10):
        query = """SELECT al.*, u.name as user_name 
                   FROM activity_log al 
                   LEFT JOIN users u ON al.user_id = u.user_id 
                   ORDER BY al.timestamp DESC 
                   LIMIT %s"""
        return DatabaseConfig.execute_query(query, (limit,), fetch=True)
    
    @staticmethod
    def get_recent_activities_by_user(user_id, limit=10):
        query = """SELECT al.*, u.name as user_name 
                   FROM activity_log al 
                   LEFT JOIN users u ON al.user_id = u.user_id 
                   WHERE al.user_id = %s
                   ORDER BY al.timestamp DESC 
                   LIMIT %s"""
        return DatabaseConfig.execute_query(query, (user_id, limit), fetch=True)

    # Session operations
    @staticmethod
    def create_session(user_id, email):
        query = "INSERT INTO sessions (user_id, email) VALUES (%s, %s)"
        return DatabaseConfig.execute_query(query, (user_id, email), lastrowid=True)

    @staticmethod
    def get_session(user_id):
        query = "SELECT * FROM sessions WHERE user_id = %s ORDER BY timestamp DESC LIMIT 1"
        return DatabaseConfig.execute_query(query, (user_id,), fetch_one=True)

# API Routes
@app.route('/api/classify-document', methods=['POST'])
def classify_document_api():
    try:
        if 'file' not in request.files:
            return jsonify({'error': 'No file provided'}), 400
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({'error': 'No file selected'}), 400
        
        with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(file.filename)[1]) as temp_file:
            file.save(temp_file.name)
            temp_path = temp_file.name
        
        try:
            extracted_text = extract_text_from_file(temp_path)
            
            if not extracted_text:
                return jsonify({
                    'documentType': 'Unclassified',
                    'confidence': 'Low',
                    'error': 'Could not extract text from file'
                }), 200
            
            document_type = classify_document(extracted_text)
            confidence = calculate_confidence(extracted_text, document_type)
            
            return jsonify({
                'documentType': document_type,
                'confidence': confidence,
                'textSample': extracted_text[:200] + '...' if len(extracted_text) > 200 else extracted_text
            })
            
        finally:
            if os.path.exists(temp_path):
                os.unlink(temp_path)
                
    except Exception as e:
        return jsonify({'error': f'Classification failed: {str(e)}'}), 500

def calculate_confidence(text, document_type):
    classification_rules = {
        "Resume / CV": ["resume", "curriculum vitae", "cv", "education", "work experience", "skills"],
        "Offer Letter / Employment Contract": ["offer of employment", "employment contract", "salary", "job title"],
        "Resignation Letter": ["resign", "resignation", "last day", "notice of resignation"],
        "Leave Request": ["leave request", "time off", "vacation", "sick leave"],
        "Employee Warning / Disciplinary Action": ["disciplinary action", "written warning", "misconduct"],
        "Performance Review": ["performance review", "evaluation", "goals", "feedback"]
    }
    
    if document_type == "Unclassified":
        return "Low"
    
    keywords = classification_rules.get(document_type, [])
    matches = sum(1 for keyword in keywords if keyword in text.lower())
    
    if matches >= 5:
        return "High"
    elif matches >= 3:
        return "Medium"
    else:
        return "Low"
def get_unique_filename_for_archive(original_name, user_id):
    """Generate unique filename when moving to archive"""
    # Removed: not used in current codebase. Keep placeholder for future archive-restore logic.
    return original_name

def get_unique_filename_for_restore(original_name, user_id):
    """Generate unique filename when restoring from archive"""
    # Removed: not used in current codebase. Keep placeholder for future archive-restore logic.
    return original_name

# Database API Routes

@app.route('/api/send-reset-email', methods=['POST'])
def send_reset_email():
    try:
        data = request.json
        email = data.get('email')
        code = data.get('code')
        
        if not email or not code:
            return jsonify({'success': False, 'error': 'Email and code are required'}), 400
        
        # Get user name for email
        user = DMSDatabase.get_user_by_email(email)
        user_name = user['name'] if user else 'User'
        
        # Send email directly using Python
        result = send_email_python(email, user_name, code)
        
        return jsonify(result)
            
    except Exception as e:
        print(f"Email sending error: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

def send_email_python(to_email, user_name, code):
    try:
        # Email configuration - USING YOUR GMAIL CREDENTIALS
        smtp_server = "smtp.gmail.com"
        smtp_port = 587
        sender_email = "plp.dmshr@gmail.com"
        sender_password = "thpn ttsr pxcw zchi"
        
        # Create message
        message = MIMEMultipart("alternative")
        message["Subject"] = "Password Reset Code - DMS Portal"
        message["From"] = "DOCUMENT MANAGEMENT SYSTEM <plp.dmshr@gmail.com>"
        message["To"] = to_email
        
        # Create HTML content
        html_content = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <style>
                body {{ font-family: Arial, sans-serif; background-color: #f4f4f4; padding: 20px; }}
                .container {{ max-width: 600px; margin: 0 auto; background: white; padding: 30px; border-radius: 10px; box-shadow: 0 2px 10px rgba(0,0,0,0.1); }}
                .header {{ text-align: center; margin-bottom: 30px; }}
                .code {{ font-size: 32px; font-weight: bold; text-align: center; color: #3b82f6; margin: 20px 0; padding: 15px; background: #f8fafc; border-radius: 8px; letter-spacing: 5px; }}
                .footer {{ margin-top: 30px; text-align: center; color: #64748b; font-size: 14px; }}
            </style>
        </head>
        <body>
            <div class='container'>
                <div class='header'>
                    <h2>Password Reset Request</h2>
                </div>
                <p>Hello {user_name},</p>
                <p>You requested a password reset for your DMS Portal account. Use the verification code below:</p>
                <div class='code'>{code}</div>
                <p>This code will expire in 10 minutes.</p>
                <p>If you didn't request this reset, please ignore this email.</p>
                <div class='footer'>
                    <p>&copy; 2025 DMS Portal. All rights reserved.</p>
                </div>
            </div>
        </body>
        </html>
        """
        
        # Create plain text content
        text_content = f"""Password Reset Code: {code}
This code will expire in 10 minutes."""
        
        # Turn these into plain/html MIMEText objects
        part1 = MIMEText(text_content, "plain")
        part2 = MIMEText(html_content, "html")
        
        # Add HTML/plain-text parts to MIMEMultipart message
        message.attach(part1)
        message.attach(part2)
        
        # Create secure connection with server and send email
        server = smtplib.SMTP(smtp_server, smtp_port)
        server.starttls()  # Secure the connection
        server.login(sender_email, sender_password)
        server.sendmail(sender_email, to_email, message.as_string())
        server.quit()
        
        print(f"✅ Email sent successfully to {to_email}")
        return {'success': True, 'message': 'Email sent successfully'}
        
    except Exception as e:
        print(f"❌ Email error: {e}")
        return {'success': False, 'message': f'Email error: {str(e)}'}

def send_file_share_email(recipient_email, recipient_name, file_name, permission_label, shared_by_name):
    """
    Send an email notification when a file is shared with a user or department.
    """
    try:
        smtp_server = "smtp.gmail.com"
        smtp_port = 587
        sender_email = "plp.dmshr@gmail.com"
        sender_password = "thpn ttsr pxcw zchi"
        
        message = MIMEMultipart("alternative")
        message["Subject"] = f"{shared_by_name} shared \"{file_name}\" with you"
        message["From"] = "DOCUMENT MANAGEMENT SYSTEM <plp.dmshr@gmail.com>"
        message["To"] = recipient_email
        
        html_content = f"""
        <!DOCTYPE html>
        <html>
          <body style="font-family: Arial, sans-serif; background-color: #f4f4f4; padding: 20px;">
            <div style="max-width: 600px; margin: 0 auto; background: white; padding: 24px; border-radius: 12px;">
              <h2 style="color:#1d4ed8;margin-top:0;">New file shared with you</h2>
              <p>Hello {recipient_name or 'colleague'},</p>
              <p><strong>{shared_by_name}</strong> shared the following file with you:</p>
              <div style="background:#f8fafc;border-left:4px solid #3b82f6;padding:12px;border-radius:8px;margin:16px 0;">
                <p style="margin:0;"><strong>File:</strong> {file_name}</p>
                <p style="margin:4px 0 0 0;"><strong>Permission:</strong> {permission_label.title()}</p>
              </div>
              <p>You can open the Document Management System to view or edit the file depending on your access.</p>
              <p style="margin-top:24px;color:#6b7280;font-size:12px;">This is an automated message from the Document Management System.</p>
            </div>
          </body>
        </html>
        """
        text_content = f"""{shared_by_name} shared "{file_name}" with you.
Permission level: {permission_label.title()}."""
        
        message.attach(MIMEText(text_content, "plain"))
        message.attach(MIMEText(html_content, "html"))
        
        server = smtplib.SMTP(smtp_server, smtp_port)
        server.starttls()
        server.login(sender_email, sender_password)
        server.sendmail(sender_email, recipient_email, message.as_string())
        server.quit()
        return True
    except Exception as e:
        print(f"Share email error: {e}")
        return False

def send_new_user_welcome_email(recipient_email, recipient_name, password):
    """
    Send a welcome email to a newly created user with their account credentials.
    """
    try:
        smtp_server = "smtp.gmail.com"
        smtp_port = 587
        sender_email = "plp.dmshr@gmail.com"
        sender_password = "thpn ttsr pxcw zchi"
        
        message = MIMEMultipart("alternative")
        message["Subject"] = "Welcome to DMS Portal - Your Account Information"
        message["From"] = "DOCUMENT MANAGEMENT SYSTEM <plp.dmshr@gmail.com>"
        message["To"] = recipient_email
        
        html_content = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <style>
                body {{
                    font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
                    background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
                    padding: 40px 20px;
                    margin: 0;
                }}
                .container {{
                    max-width: 600px;
                    margin: 0 auto;
                    background: white;
                    border-radius: 16px;
                    overflow: hidden;
                    box-shadow: 0 20px 60px rgba(0,0,0,0.3);
                }}
                .header {{
                    background: linear-gradient(135deg, #3b82f6 0%, #2563eb 100%);
                    padding: 40px 30px;
                    text-align: center;
                    color: white;
                }}
                .header h1 {{
                    margin: 0;
                    font-size: 28px;
                    font-weight: 600;
                }}
                .header p {{
                    margin: 10px 0 0 0;
                    opacity: 0.9;
                    font-size: 16px;
                }}
                .content {{
                    padding: 40px 30px;
                }}
                .welcome-text {{
                    color: #334155;
                    font-size: 16px;
                    line-height: 1.6;
                    margin-bottom: 30px;
                }}
                .credentials-box {{
                    background: #f8fafc;
                    border-left: 4px solid #3b82f6;
                    padding: 25px;
                    border-radius: 8px;
                    margin: 25px 0;
                }}
                .credential-row {{
                    display: flex;
                    margin: 15px 0;
                    align-items: center;
                }}
                .credential-label {{
                    font-weight: 600;
                    color: #64748b;
                    min-width: 100px;
                    font-size: 14px;
                }}
                .credential-value {{
                    color: #1e293b;
                    font-family: 'Courier New', monospace;
                    font-size: 15px;
                    font-weight: 600;
                    background: white;
                    padding: 8px 12px;
                    border-radius: 6px;
                    border: 1px solid #e2e8f0;
                }}
                .password-value {{
                    color: #3b82f6;
                    font-size: 20px;
                    letter-spacing: 3px;
                }}
                .important-note {{
                    background: #fef3c7;
                    border-left: 4px solid #f59e0b;
                    padding: 20px;
                    border-radius: 8px;
                    margin: 25px 0;
                }}
                .important-note p {{
                    margin: 0;
                    color: #78350f;
                    font-size: 14px;
                    line-height: 1.6;
                }}
                .footer {{
                    background: #f8fafc;
                    padding: 30px;
                    text-align: center;
                    border-top: 1px solid #e2e8f0;
                }}
                .footer p {{
                    margin: 5px 0;
                    color: #64748b;
                    font-size: 13px;
                }}
                .icon {{
                    display: inline-block;
                    font-size: 48px;
                    margin-bottom: 10px;
                }}
            </style>
        </head>
        <body>
            <div class='container'>
                <div class='header'>
                    <div class='icon'>🎉</div>
                    <h1>Welcome to DMS Portal!</h1>
                    <p>Your account has been successfully created</p>
                </div>
                
                <div class='content'>
                    <div class='welcome-text'>
                        <p>Hello <strong>{recipient_name}</strong>,</p>
                        <p>Your account has been created in the Document Management System. You can now access the portal using the credentials below:</p>
                    </div>
                    
                    <div class='credentials-box'>
                        <div class='credential-row'>
                            <span class='credential-label'>📧 Email:</span>
                            <span class='credential-value'>{recipient_email}</span>
                        </div>
                        <div class='credential-row'>
                            <span class='credential-label'>🔑 Password:</span>
                            <span class='credential-value password-value'>{password}</span>
                        </div>
                    </div>
                    
                    <div class='important-note'>
                        <p><strong>⚠️ Important Security Note:</strong></p>
                        <p>This is a temporary password. For security reasons, we strongly recommend changing your password after your first login. You can update your password in your profile settings.</p>
                    </div>
                </div>
                
                <div class='footer'>
                    <p><strong>Document Management System</strong></p>
                    <p>&copy; 2025 DMS Portal. All rights reserved.</p>
                    <p style='margin-top: 15px; font-size: 12px;'>This is an automated message. Please do not reply to this email.</p>
                </div>
            </div>
        </body>
        </html>
        """
        
        text_content = f"""Welcome to DMS Portal!

Hello {recipient_name},

Your account has been created. Here are your login credentials:

Email: {recipient_email}
Password: {password}

IMPORTANT: Please change your password after your first login for security.

Best regards,
DMS Portal Team
"""
        
        part1 = MIMEText(text_content, "plain")
        part2 = MIMEText(html_content, "html")
        message.attach(part1)
        message.attach(part2)
        
        server = smtplib.SMTP(smtp_server, smtp_port)
        server.starttls()
        server.login(sender_email, sender_password)
        server.sendmail(sender_email, recipient_email, message.as_string())
        server.quit()
        
        print(f"Succesfully added user. Email sent")
        return {'success': True, 'message': 'Welcome email sent successfully'}
        
    except Exception as e:
        print(f"❌ Email error: {e}")
        return {'success': False, 'message': f'Email error: {str(e)}'}

@app.route('/api/debug/user/<int:user_id>', methods=['GET'])
def debug_user(user_id):
    """Debug endpoint to check user data"""
    user = DMSDatabase.get_user_by_id(user_id)
    if user:
        user_serialized = serialize_user_record(user)
        return jsonify({
            'success': True,
            'user': {
                'user_id': user_serialized['user_id'],
                'name': user_serialized['name'],
                'email': user_serialized['email'],
                'password': user_serialized.get('password'),  # This will show current password
                'role': user_serialized['role'],
                'status': user_serialized['status'],
                'profile_image': user_serialized.get('profile_image')
            }
        })
    else:
        return jsonify({'success': False, 'error': 'User not found'})
    
@app.route('/api/users', methods=['GET'])
def get_users():
    users = DMSDatabase.get_all_users() or []
    serialized = [serialize_user_record(user) for user in users]
    return jsonify(serialized)

@app.route('/api/users', methods=['POST'])
def create_user():
    import random
    import string
    
    data = request.json
    
    # Generate a random 6-character password (mix of letters and numbers)
    password = ''.join(random.choices(string.ascii_letters + string.digits, k=6))
    
    user_id = DMSDatabase.create_user(
        data['name'], data['email'], password, 
        data['role'], data['department_id']
    )
    
    # Send welcome email with credentials
    if user_id:
        email_result = send_new_user_welcome_email(
            data['email'], 
            data['name'], 
            password
        )
        
        if email_result.get('success'):
            print(f"✅ User created and welcome email sent to {data['email']}")
        else:
            print(f"⚠️ User created but email failed: {email_result.get('message')}")
    
    return jsonify({'user_id': user_id, 'success': True})

@app.route('/api/users/<int:user_id>/verify-password', methods=['POST'])
def verify_user_password(user_id):
    """Verify if the provided current password matches the user's password"""
    try:
        data = request.json
        current_password = data.get('current_password')
        
        if not current_password:
            return jsonify({'success': False, 'error': 'Current password is required'}), 400
        
        # Get user from database
        user = DMSDatabase.get_user_by_id(user_id)
        if not user:
            return jsonify({'success': False, 'error': 'User not found'}), 404
        
        # Check if password matches
        if user.get('password') == current_password:
            return jsonify({'success': True, 'message': 'Password verified'})
        else:
            return jsonify({'success': False, 'error': 'Incorrect current password'}), 401
            
    except Exception as e:
        print(f"Error verifying password for user {user_id}: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({'success': False, 'error': 'Failed to verify password'}), 500

@app.route('/api/users/<int:user_id>', methods=['PUT'])
def update_user(user_id):
    try:
        data = request.json
        print(f"DEBUG: Updating user {user_id} with data: {data}")
        
        # Check if password is being updated
        password = data.get('password')
        
        profile_image_token = normalize_profile_image_token(data.get('profile_image'))

        # Determine if department_id is explicitly intended to be cleared
        clear_department = 'department_id' in data and data.get('department_id') is None

        result = DMSDatabase.update_user(
            user_id, 
            data.get('name'), 
            data.get('email'), 
            data.get('role'),
            data.get('department_id'), 
            data.get('status'), 
            profile_image_token,
            password,  # Pass password if provided
            clear_department=clear_department
        )
        
        if result:
            return jsonify({'success': True, 'message': 'User updated successfully'})
        else:
            return jsonify({'success': False, 'error': 'Failed to update user'}), 500
            
    except Exception as e:
        print(f"Error in update_user API: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({'success': False, 'error': f'Failed to update user: {str(e)}'}), 500
    

@app.route('/api/users/<int:user_id>/archive', methods=['POST'])
def archive_user(user_id):
    try:
        data = request.json or {}
        status = data.get('status', 'archived')
        
        # Update user status to archived using update_user with only status parameter
        result = DMSDatabase.update_user(user_id, None, None, None, None, status, None, None)
        if result:
            return jsonify({'success': True, 'message': 'User archived successfully'})
        else:
            return jsonify({'success': False, 'error': 'Failed to archive user'}), 500
    except Exception as e:
        print(f"Error archiving user {user_id}: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({'success': False, 'error': f'Failed to archive user: {str(e)}'}), 500

@app.route('/api/upload-profile-image', methods=['POST'])
def upload_profile_image():
    try:
        if 'image' not in request.files:
            return jsonify({'error': 'No image provided'}), 400
        
        file = request.files['image']
        user_id = request.form.get('user_id')
        
        if not user_id:
            return jsonify({'error': 'User ID is required'}), 400
        
        if file.filename == '':
            return jsonify({'error': 'Please select an image file'}), 400
        
        # Validate image type
        if not file.content_type.startswith('image/'):
            return jsonify({'error': 'File must be an image'}), 400
        
        # Check file size (max 5MB)
        file.seek(0, 2)
        file_size = file.tell()
        file.seek(0)
        
        if file_size > 5 * 1024 * 1024:
            return jsonify({'error': 'Image size must be less than 5MB'}), 400
        
        # Save image to uploads/profile_images directory
        upload_folder = os.path.join(os.path.dirname(__file__), 'uploads', 'profile_images')
        os.makedirs(upload_folder, exist_ok=True)
        
        # Generate unique filename
        import time
        file_ext = file.filename.rsplit('.', 1)[1].lower() if '.' in file.filename else 'jpg'
        filename = f"profile_{user_id}_{int(time.time())}.{file_ext}"
        file_path = os.path.join(upload_folder, filename)
        
        file.save(file_path)
        
        # Hash filename for storage and rename file to hashed version
        image_hash = hash_file_path(file_path)
        hashed_filename = f"{image_hash}.{file_ext}"
        hashed_path = os.path.join(upload_folder, hashed_filename)
        try:
            if os.path.exists(hashed_path):
                os.remove(hashed_path)
            os.rename(file_path, hashed_path)
            stored_filename = hashed_filename
        except Exception as rename_error:
            print(f"Profile image rename error: {rename_error}")
            stored_filename = filename
            hashed_path = file_path
        
        image_token = stored_filename
        image_url = build_profile_image_url(image_token)
        
        # Update user profile_image in database using hashed token
        DMSDatabase.update_user(user_id, None, None, None, None, None, image_token, None)
        
        return jsonify({'success': True, 'image_url': image_url, 'image_token': image_token})
        
    except Exception as e:
        print(f"Error uploading profile image: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': f'Failed to upload profile image: {str(e)}'}), 500

@app.route('/api/uploads/profile_images/<filename>')
def serve_profile_image(filename):
    """Serve profile images"""
    try:
        image_path = os.path.join(os.path.dirname(__file__), 'uploads', 'profile_images', filename)
        if os.path.exists(image_path):
            return send_file(image_path)
        else:
            return jsonify({'error': 'Image not found'}), 404
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/uploads/company_logo/<filename>')
def serve_company_logo(filename):
    """Serve company logo images"""
    try:
        logo_path = os.path.join(os.path.dirname(__file__), 'uploads', 'company_logo', filename)
        if os.path.exists(logo_path):
            return send_file(logo_path)
        else:
            return jsonify({'error': 'Logo not found'}), 404
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/users/<int:user_id>/restore', methods=['POST'])
def restore_user(user_id):
    try:
        # Update user status to active
        result = DMSDatabase.update_user(user_id, None, None, None, None, 'active', None, None)
        if result:
            return jsonify({'success': True, 'message': 'User restored successfully'})
        else:
            return jsonify({'success': False, 'error': 'Failed to restore user'}), 500
    except Exception as e:
        return jsonify({'success': False, 'error': f'Failed to restore user: {str(e)}'}), 500

@app.route('/api/files', methods=['GET'])
def get_files():
    query = """SELECT f.*, u.name as owner_name, c.name as category_name 
               FROM files f 
               LEFT JOIN users u ON f.user_id = u.user_id 
               LEFT JOIN categories c ON f.category_id = c.category_id"""
    files = DatabaseConfig.execute_query(query, fetch=True) or []
    
    request_user = get_request_user()
    if not request_user:
        return jsonify({'error': 'Unauthorized'}), 401
    
    if request_user.get('role') == 'system_admin':
        annotated = []
        for file in files:
            file_copy = dict(file)
            if request_user:
                if file_copy['user_id'] == request_user['user_id']:
                    permission = 'owner'
                else:
                    permission = 'admin'
            else:
                permission = None
            file_copy['current_user_permission'] = permission
            file_copy['can_edit'] = permission in ('owner', 'admin')
            annotated.append(file_copy)
        return jsonify(annotated)
    
    user_id = request_user['user_id']
    dept_id = request_user.get('department_id')
    
    user_shares = DatabaseConfig.execute_query(
        "SELECT file_id, permission FROM file_shares WHERE shared_with_user_id = %s",
        (user_id,),
        fetch=True
    ) or []
    dept_shares = []
    if dept_id:
        dept_shares = DatabaseConfig.execute_query(
            "SELECT file_id, permission FROM file_shares WHERE shared_with_department_id = %s",
            (dept_id,),
            fetch=True
        ) or []
    
    user_share_map = {}
    for row in user_shares:
        user_share_map[row['file_id']] = merge_permission(user_share_map.get(row['file_id']), row.get('permission', 'viewer'))
    
    dept_share_map = {}
    for row in dept_shares:
        dept_share_map[row['file_id']] = merge_permission(dept_share_map.get(row['file_id']), row.get('permission', 'viewer'))
    
    allowed_files = []
    for file in files:
        permission = None
        if file['user_id'] == user_id:
            permission = 'owner'
        elif user_share_map.get(file['file_id']):
            permission = user_share_map[file['file_id']]
        elif dept_id and dept_share_map.get(file['file_id']):
            permission = dept_share_map[file['file_id']]
        
        if permission:
            file_copy = dict(file)
            file_copy['current_user_permission'] = permission
            file_copy['can_edit'] = permission in ('owner', 'editor')
            allowed_files.append(file_copy)
    
    return jsonify(allowed_files)

@app.route('/api/all-files', methods=['GET'])
def get_all_files_including_archived():
    query = """SELECT f.*, u.name as owner_name, c.name as category_name 
               FROM files f 
               LEFT JOIN users u ON f.user_id = u.user_id 
               LEFT JOIN categories c ON f.category_id = c.category_id"""
    files = DatabaseConfig.execute_query(query, fetch=True)
    return jsonify(files or [])

@app.route('/api/files', methods=['POST'])
def create_file():
    data = request.json
    file_id = DMSDatabase.create_file(data)
    return jsonify({'file_id': file_id})

@app.route('/api/files/<int:file_id>/shares', methods=['GET'])
def get_file_shares(file_id):
    try:
        query = """SELECT fs.*, 
                   u.name as shared_with_user_name, 
                   d.name as shared_with_department_name
                   FROM file_shares fs
                   LEFT JOIN users u ON fs.shared_with_user_id = u.user_id
                   LEFT JOIN departments d ON fs.shared_with_department_id = d.department_id
                   WHERE fs.file_id = %s"""
        shares = DatabaseConfig.execute_query(query, (file_id,), fetch=True)
        return jsonify(shares or [])
    except Exception as e:
        print(f"Error getting file shares: {str(e)}")
        return jsonify({'error': f'Failed to get file shares: {str(e)}'}), 500

@app.route('/api/files/<int:file_id>/share', methods=['POST'])
def share_file(file_id):
    try:
        data = request.json
        shared_with_user_id = data.get('shared_with_user_id')
        shared_with_department_id = data.get('shared_with_department_id')
        permission = data.get('permission', 'viewer')
        file_record = DMSDatabase.get_file_by_id(file_id)
        if not file_record:
            return jsonify({'success': False, 'error': 'File not found'}), 404
        
        request_user = get_request_user()
        if not request_user:
            return jsonify({'success': False, 'error': 'Unauthorized'}), 401
        
        request_permission = get_user_file_permission(file_record, request_user)
        if request_permission not in ('owner', 'admin'):
            return jsonify({'success': False, 'error': 'You do not have permission to share this file.'}), 403
        
        shared_by = data.get('shared_by') or request_user['user_id']
        shared_by_user = DMSDatabase.get_user_by_id(shared_by) or {}
        
        if not shared_with_user_id and not shared_with_department_id:
            return jsonify({'success': False, 'error': 'Must specify user or department'}), 400
        
        # Check if share already exists
        check_query = """SELECT share_id FROM file_shares 
                        WHERE file_id = %s 
                        AND (shared_with_user_id = %s OR shared_with_department_id = %s)"""
        existing = DatabaseConfig.execute_query(
            check_query, 
            (file_id, shared_with_user_id or 0, shared_with_department_id or 0), 
            fetch=True
        )
        
        if existing:
            return jsonify({'success': False, 'error': 'Already shared with this user/department'}), 400
        
        # Create share
        insert_query = """INSERT INTO file_shares 
                         (file_id, shared_with_user_id, shared_with_department_id, permission, shared_by)
                         VALUES (%s, %s, %s, %s, %s)"""
        DatabaseConfig.execute_query(
            insert_query,
            (file_id, shared_with_user_id, shared_with_department_id, permission, shared_by)
        )
        
        # Update file shared flag
        DMSDatabase.update_file(file_id, {'shared': True})
        
        # Send notification emails
        notify_targets = []
        if shared_with_user_id:
            target_user = DMSDatabase.get_user_by_id(shared_with_user_id)
            if target_user and target_user.get('email'):
                notify_targets.append(target_user)
        elif shared_with_department_id:
            dept_users = DatabaseConfig.execute_query(
                "SELECT name, email FROM users WHERE department_id = %s AND status = 'active'",
                (shared_with_department_id,),
                fetch=True
            ) or []
            notify_targets.extend([u for u in dept_users if u.get('email')])
        
        file_name = file_record.get('name') or file_record.get('original_name') or f"File #{file_id}"
        shared_by_name = shared_by_user.get('name', 'A colleague')
        for recipient in notify_targets:
            send_file_share_email(recipient.get('email'), recipient.get('name'), file_name, permission, shared_by_name)
        
        return jsonify({'success': True, 'message': 'File shared successfully'})
    except Exception as e:
        print(f"Error sharing file: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({'success': False, 'error': f'Failed to share file: {str(e)}'}), 500

@app.route('/api/file-shares/<int:share_id>', methods=['PUT'])
def update_file_share(share_id):
    try:
        data = request.json
        permission = data.get('permission')
        
        if permission not in ['viewer', 'editor']:
            return jsonify({'success': False, 'error': 'Invalid permission'}), 400
        
        query = "UPDATE file_shares SET permission = %s WHERE share_id = %s"
        DatabaseConfig.execute_query(query, (permission, share_id))
        
        return jsonify({'success': True, 'message': 'Share updated successfully'})
    except Exception as e:
        print(f"Error updating file share: {str(e)}")
        return jsonify({'success': False, 'error': f'Failed to update share: {str(e)}'}), 500

@app.route('/api/file-shares/<int:share_id>', methods=['DELETE'])
def delete_file_share(share_id):
    try:
        # Get file_id before deleting
        file_id_query = "SELECT file_id FROM file_shares WHERE share_id = %s"
        file_result = DatabaseConfig.execute_query(file_id_query, (share_id,), fetch=True)
        
        query = "DELETE FROM file_shares WHERE share_id = %s"
        DatabaseConfig.execute_query(query, (share_id,))
        
        # Check if file still has shares
        if file_result:
            file_id = file_result[0]['file_id']
            check_query = "SELECT COUNT(*) as count FROM file_shares WHERE file_id = %s"
            result = DatabaseConfig.execute_query(check_query, (file_id,), fetch=True)
            
            if result and result[0]['count'] == 0:
                # No more shares, set file to not shared
                DMSDatabase.update_file(file_id, {'shared': False})
        
        return jsonify({'success': True, 'message': 'Share removed successfully'})
    except Exception as e:
        print(f"Error deleting file share: {str(e)}")
        return jsonify({'success': False, 'error': f'Failed to remove share: {str(e)}'}), 500

@app.route('/api/files/<int:file_id>', methods=['PUT'])
def update_file(file_id):
    try:
        data = request.json
        result = DMSDatabase.update_file(file_id, data)
        return jsonify({'success': True})
    except Exception as e:
        return jsonify({'error': f'Failed to update file: {str(e)}'}), 500

@app.route('/api/files/<int:file_id>/rename', methods=['POST'])
def rename_file(file_id):
    try:
        data = request.json
        new_name = data.get('new_name')
        
        if not new_name:
            return jsonify({'success': False, 'error': 'New name is required'}), 400
        
        # Get current file record
        file_record = DMSDatabase.get_file_by_id(file_id)
        if not file_record:
            return jsonify({'success': False, 'error': 'File not found'}), 404

        # Permission: only owner/editor/admin can rename
        request_user = get_request_user()
        if not request_user:
            return jsonify({'success': False, 'error': 'Unauthorized'}), 401
        permission = get_user_file_permission(file_record, request_user)
        if permission not in ('owner', 'editor', 'admin'):
            return jsonify({'success': False, 'error': 'Access denied'}), 403
        
        user_id = file_record['user_id']
        original_name = file_record['name']
        
        # Check if new name already exists for this user (excluding current file)
        query = "SELECT file_id FROM files WHERE user_id = %s AND status = 'active' AND name = %s AND file_id != %s"
        connection = DatabaseConfig.get_connection()
        cursor = connection.cursor(dictionary=True)
        cursor.execute(query, (user_id, new_name, file_id))
        existing_file = cursor.fetchone()
        cursor.close()
        connection.close()
        
        if existing_file:
            # File exists - we need to overwrite it
            # Delete the existing file first
            existing_file_record = DMSDatabase.get_file_by_id(existing_file['file_id'])
            if existing_file_record:
                # Find and delete the old file's physical file
                old_file_path_hash = existing_file_record.get('file_path')
                old_file_path = None
                
                # Try to find the actual file path
                upload_folder = app.config['UPLOAD_FOLDER']
                for root, dirs, files in os.walk(upload_folder):
                    for file in files:
                        if file == existing_file_record['name'] or file == existing_file_record.get('original_name'):
                            old_file_path = os.path.join(root, file)
                            break
                    if old_file_path:
                        break
                
                if old_file_path and os.path.exists(old_file_path):
                    try:
                        os.remove(old_file_path)
                    except Exception as e:
                        print(f"Warning: Could not delete old file: {e}")
                
                # Delete the old file record
                DMSDatabase.delete_file(existing_file['file_id'])
        
        # Find the actual file path for the current file
        file_path_hash = file_record.get('file_path')
        actual_file_path = None
        
        # Try to find the actual file path
        upload_folder = app.config['UPLOAD_FOLDER']
        for root, dirs, files in os.walk(upload_folder):
            for file in files:
                if file == original_name or file == file_record.get('original_name'):
                    actual_file_path = os.path.join(root, file)
                    break
            if actual_file_path:
                break
        
        # If we found the file, rename it physically
        if actual_file_path and os.path.exists(actual_file_path):
            try:
                # Get the directory and new file path
                file_dir = os.path.dirname(actual_file_path)
                new_file_path = os.path.join(file_dir, secure_filename(new_name))
                
                # Rename the physical file
                os.rename(actual_file_path, new_file_path)
                
                # Update the file_path hash in database
                new_file_path_hash = hash_file_path(new_file_path)
                DMSDatabase.update_file(file_id, {
                    'name': new_name,
                    'file_path': new_file_path_hash,
                    'is_encrypted': bool(new_file_path_hash),
                    'encryption_version': current_encryption_version()
                })
                
                print(f"DEBUG: File renamed successfully: {original_name} -> {new_name}")
                return jsonify({'success': True, 'message': 'File renamed successfully'})
                
            except Exception as e:
                print(f"Warning: Could not rename physical file: {e}")
                # Still update the name in database even if physical rename fails
                DMSDatabase.update_file(file_id, {'name': new_name})
                
                print(f"DEBUG: File renamed in database only: {original_name} -> {new_name}")
                return jsonify({'success': True, 'message': 'File renamed successfully (database only)'})
        else:
            # File not found physically, just update the name in database
            DMSDatabase.update_file(file_id, {'name': new_name})
            
            print(f"DEBUG: File renamed in database (no physical file): {original_name} -> {new_name}")
        
        return jsonify({'success': True, 'message': 'File renamed successfully'})
            
    except Exception as e:
        print(f"Error renaming file: {e}")
        return jsonify({'success': False, 'error': f'Failed to rename file: {str(e)}'}), 500

@app.route('/api/files/<int:file_id>/archive', methods=['POST'])
def archive_file_api(file_id):
    try:
        print(f"DEBUG: API - Archiving file {file_id}")
        
        # First check if file exists
        file_record = DMSDatabase.get_file_by_id(file_id)
        if not file_record:
            print(f"DEBUG: API - File {file_id} not found")
            return jsonify({'success': False, 'error': 'File not found'}), 404

        # Permission: only owner/editor/admin can archive
        request_user = get_request_user()
        if not request_user:
          return jsonify({'success': False, 'error': 'Unauthorized'}), 401
        permission = get_user_file_permission(file_record, request_user)
        if permission not in ('owner', 'editor', 'admin'):
          return jsonify({'success': False, 'error': 'Access denied'}), 403
        
        print(f"DEBUG: API - File found: {file_record['name']}")
        
        # Archive the file using the database method
        result = DMSDatabase.archive_file(file_id)
        
        if result:
            print(f"DEBUG: API - File {file_id} archived successfully")
            return jsonify({'success': True, 'message': 'File archived successfully'})
        else:
            print(f"DEBUG: API - Database method returned false for file {file_id}")
            # Double check if the file was actually archived
            updated_file = DMSDatabase.get_file_by_id(file_id)
            if updated_file and updated_file.get('status') == 'archived':
                print(f"DEBUG: API - File {file_id} is archived despite method returning false")
                return jsonify({'success': True, 'message': 'File archived successfully'})
            else:
                return jsonify({'success': False, 'error': 'Failed to archive file'}), 500
            
    except Exception as e:
        print(f"API Archive error for file {file_id}: {str(e)}")
        return jsonify({'success': False, 'error': f'Archive failed: {str(e)}'}), 500
    
@app.route('/api/files/<int:file_id>/update-workspaces', methods=['POST'])
def update_file_workspaces_api(file_id):
    """Update all workspace relationships for a file"""
    try:
        data = request.json
        workspace_ids = data.get('workspace_ids', [])

        file_record = DMSDatabase.get_file_by_id(file_id)
        if not file_record:
            return jsonify({'error': 'File not found'}), 404

        # Permission: only owner/editor/admin can manage workspaces
        request_user = get_request_user()
        if not request_user:
            return jsonify({'error': 'Unauthorized'}), 401
        permission = get_user_file_permission(file_record, request_user)
        if permission not in ('owner', 'editor', 'admin'):
            return jsonify({'error': 'Access denied'}), 403
        
        print(f"DEBUG: API - Updating workspaces for file {file_id} to: {workspace_ids}")
        
        # Convert all workspace IDs to integers
        workspace_ids = [int(wid) for wid in workspace_ids]
        
        result = DMSDatabase.update_file_workspaces(file_id, workspace_ids)
        
        if result:
            return jsonify({
                'success': True, 
                'message': f'File updated in {len(workspace_ids)} workspace(s)',
                'workspace_ids': workspace_ids
            })
        else:
            return jsonify({'error': 'Failed to update file workspaces'}), 500
            
    except Exception as e:
        print(f"Error in update_file_workspaces_api: {e}")
        return jsonify({'error': f'Failed to update file workspaces: {str(e)}'}), 500
    
@app.route('/api/debug/file-status/<int:file_id>', methods=['GET'])
def debug_file_status(file_id):
    """Debug endpoint to check file status"""
    file_record = DMSDatabase.get_file_by_id(file_id)
    if file_record:
        return jsonify({
            'success': True,
            'file_id': file_id,
            'name': file_record.get('name'),
            'original_name': file_record.get('original_name'),
            'status': file_record.get('status', 'unknown'),
            'user_id': file_record.get('user_id')
        })
    else:
        return jsonify({'success': False, 'error': 'File not found'})

@app.route('/api/files/<int:file_id>', methods=['DELETE'])
def delete_file(file_id):
    try:
        file_record = DMSDatabase.get_file_by_id(file_id)
        if not file_record:
            return jsonify({'success': False, 'error': 'File not found'}), 404

        request_user = get_request_user()
        if not request_user:
            return jsonify({'success': False, 'error': 'Unauthorized'}), 401

        permission = get_user_file_permission(file_record, request_user)
        if permission not in ('owner', 'admin'):
            return jsonify({'success': False, 'error': 'Access denied'}), 403

        result = DMSDatabase.delete_file(file_id)
        
        if result:
            print(f"DEBUG: File {file_id} deleted successfully")
            
            if request_user:
                DMSDatabase.log_activity(
                    request_user['user_id'], 
                    f"Permanently deleted file: {file_record.get('name')}"
                )
            
            return jsonify({'success': True, 'message': 'File deleted permanently'})
        else:
            print(f"DEBUG: Failed to delete file {file_id}")
            return jsonify({'success': False, 'error': 'Failed to delete file'}), 500
            
    except Exception as e:
        print(f"Delete file API error: {e}")
        return jsonify({'success': False, 'error': f'Delete failed: {str(e)}'}), 500

@app.route('/api/files/<int:file_id>/restore', methods=['POST'])
def restore_file(file_id):
    try:
        print(f"DEBUG: Restoring file {file_id}")
        
        # First get the file details
        file_record = DMSDatabase.get_file_by_id(file_id)
        if not file_record:
            print(f"DEBUG: File {file_id} not found")
            return jsonify({'success': False, 'error': 'File not found'}), 404
        
        # Generate unique name for restore
        original_name = file_record['original_name']
        user_id = file_record['user_id']
        
        # Get unique filename for active files
        base_name = original_name.rsplit('.', 1)[0]
        extension = original_name.rsplit('.', 1)[1].lower()
        
        query = "SELECT name FROM files WHERE user_id = %s AND status = 'active' AND name LIKE %s"
        connection = DatabaseConfig.get_connection()
        cursor = connection.cursor(dictionary=True)
        
        like_pattern = f"{base_name}%.{extension}"
        cursor.execute(query, (user_id, like_pattern))
        existing_files = cursor.fetchall()
        cursor.close()
        connection.close()
        
        final_filename = original_name
        counter = 1
        
        existing_names = [f['name'] for f in existing_files]
        while final_filename in existing_names:
            final_filename = f"{base_name} ({counter}).{extension}"
            counter += 1
        
        print(f"DEBUG: Updating file {file_id} to active with name: {final_filename}")
        
        # Update status and name for restore - USING DIRECT DATABASE CONNECTION
        query = "UPDATE files SET status='active', name=%s WHERE file_id=%s"
        connection = DatabaseConfig.get_connection()
        cursor = connection.cursor()
        cursor.execute(query, (final_filename, file_id))
        affected_rows = cursor.rowcount
        connection.commit()
        cursor.close()
        connection.close()
        
        print(f"DEBUG: Restore query affected {affected_rows} rows")
        
        if affected_rows > 0:
            print(f"DEBUG: File {file_id} restored successfully as {final_filename}")
            return jsonify({'success': True, 'message': 'File restored successfully', 'new_name': final_filename})
        else:
            print(f"DEBUG: No rows affected when restoring file {file_id}")
            return jsonify({'success': False, 'error': 'Failed to restore file'}), 500
            
    except Exception as e:
        print(f"Restore error for file {file_id}: {e}")
        return jsonify({'success': False, 'error': f'Restore failed: {str(e)}'}), 500

@app.route('/api/debug/file/<int:file_id>', methods=['GET'])
def debug_file(file_id):
    """Debug endpoint to check file status"""
    file_record = DMSDatabase.get_file_by_id(file_id)
    if file_record:
        return jsonify({
            'success': True,
            'file': file_record,
            'status': file_record.get('status', 'unknown')
        })
    else:
        return jsonify({'success': False, 'error': 'File not found'})
    
@app.route('/api/debug/workspace/<int:workspace_id>/files-count', methods=['GET'])
def debug_workspace_files_count(workspace_id):
    """Debug endpoint to check files in workspace"""
    try:
        # Check old way
        old_query = "SELECT COUNT(*) as count FROM files WHERE workspace_id = %s AND status = 'active'"
        old_count = DatabaseConfig.execute_query(old_query, (workspace_id,), fetch_one=True)
        
        # Check new way
        new_query = "SELECT COUNT(*) as count FROM file_workspaces WHERE workspace_id = %s"
        new_count = DatabaseConfig.execute_query(new_query, (workspace_id,), fetch_one=True)
        
        # Get actual files
        files = DMSDatabase.get_files_by_workspace(workspace_id)
        
        return jsonify({
            'workspace_id': workspace_id,
            'old_way_count': old_count['count'] if old_count else 0,
            'new_way_count': new_count['count'] if new_count else 0,
            'actual_files_count': len(files),
            'files': [f['file_id'] for f in files]
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/debug/file/<int:file_id>/workspaces', methods=['GET'])
def debug_file_workspaces(file_id):
    """Debug endpoint to check which workspaces a file is in"""
    try:
        # Check file_workspaces table
        query = "SELECT workspace_id FROM file_workspaces WHERE file_id = %s"
        workspaces = DatabaseConfig.execute_query(query, (file_id,), fetch=True) or []
        
        # Check files table
        file_record = DMSDatabase.get_file_by_id(file_id)
        
        return jsonify({
            'file_id': file_id,
            'in_file_workspaces': [w['workspace_id'] for w in workspaces],
            'files_table_workspace_id': file_record.get('workspace_id') if file_record else None,
            'file_status': file_record.get('status') if file_record else None
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/files/<int:file_id>/workspaces', methods=['GET'])
def get_file_workspaces(file_id):
    """Get all workspaces where a file exists"""
    try:
        request_user = get_request_user()
        if not request_user:
            return jsonify({'error': 'Unauthorized'}), 401
        file_record = DMSDatabase.get_file_by_id(file_id)
        if not file_record:
            return jsonify({'error': 'File not found'}), 404
        permission = get_user_file_permission(file_record, request_user)
        if permission not in ('owner', 'editor', 'admin'):
            return jsonify({'error': 'Access denied'}), 403
        workspaces = DMSDatabase.get_file_workspaces(file_id)
        return jsonify(workspaces)
    except Exception as e:
        return jsonify({'error': f'Failed to get file workspaces: {str(e)}'}), 500

@app.route('/api/files/<int:file_id>/remove-from-workspace', methods=['POST'])
def remove_file_from_workspace_api(file_id):
    """Remove file from a specific workspace"""
    try:
        data = request.json
        workspace_id = data.get('workspace_id')
        
        if not workspace_id:
            return jsonify({'error': 'Workspace ID is required'}), 400
        
        result = DMSDatabase.remove_file_from_workspace(file_id, workspace_id)
        
        if result:
            return jsonify({'success': True, 'message': 'File removed from workspace'})
        else:
            return jsonify({'error': 'Failed to remove file from workspace'}), 500
            
    except Exception as e:
        return jsonify({'error': f'Failed to remove file from workspace: {str(e)}'}), 500

@app.route('/api/files/<int:file_id>/check-workspace/<int:workspace_id>', methods=['GET'])
def check_file_in_workspace(file_id, workspace_id):
    """Check if file is in a specific workspace"""
    try:
        is_in_workspace = DMSDatabase.is_file_in_workspace(file_id, workspace_id)
        return jsonify({'in_workspace': is_in_workspace})
    except Exception as e:
        return jsonify({'error': f'Failed to check file in workspace: {str(e)}'}), 500
    
@app.route('/api/departments', methods=['GET'])
def get_departments():
    departments = DMSDatabase.get_all_departments()
    return jsonify(departments or [])

@app.route('/api/departments', methods=['POST'])
def create_department():
    data = request.json
    department_id = DMSDatabase.create_department(data['name'], data['user_id'])
    return jsonify({'department_id': department_id})

@app.route('/api/departments/<int:department_id>', methods=['DELETE'])
def delete_department(department_id):
    try:
        connection = DatabaseConfig.get_connection()
        cursor = connection.cursor()

        # Ensure an 'Unassigned' department exists
        cursor.execute("SELECT department_id FROM departments WHERE name=%s", ("Unassigned",))
        row = cursor.fetchone()
        if row:
            unassigned_id = row[0]
        else:
            cursor.execute("SELECT COALESCE(MAX(department_id), 4000) + 1 FROM departments")
            unassigned_id = cursor.fetchone()[0]
            cursor.execute(
                "INSERT INTO departments (department_id, name, user_id) VALUES (%s, %s, %s)",
                (unassigned_id, "Unassigned", 1)
            )
            print(f"DEBUG: Created 'Unassigned' department id={unassigned_id}")

        # Reassign members to 'Unassigned'
        cursor.execute(
            "UPDATE users SET department_id=%s WHERE department_id=%s",
            (unassigned_id, department_id)
        )
        reassigned = cursor.rowcount

        # Attempt to delete the department
        cursor.execute("DELETE FROM departments WHERE department_id=%s", (department_id,))
        deleted_rows = cursor.rowcount

        connection.commit()
        cursor.close()
        connection.close()

        if deleted_rows > 0:
            return jsonify({'success': True, 'reassigned_members': reassigned})
        else:
            return jsonify({'success': False, 'error': 'Department not found'}), 404
    except Exception as e:
        print(f"Error deleting department {department_id}: {e}")
        try:
            if 'connection' in locals():
                connection.rollback()
        except Exception:
            pass
        import traceback
        traceback.print_exc()
        return jsonify({'success': False, 'error': f'Failed to delete department: {str(e)}'}), 500

@app.route('/api/categories', methods=['GET'])
def get_categories():
    categories = DMSDatabase.get_all_categories()
    return jsonify(categories or [])

@app.route('/api/categories', methods=['POST'])
def create_category():
    data = request.json
    category_id = DMSDatabase.create_category(
        data['name'], data.get('description'), 
        data.get('auto_created', False), data.get('is_unclassified', False)
    )
    return jsonify({'category_id': category_id})

@app.route('/api/categories/<int:category_id>', methods=['PUT'])
def update_category(category_id):
    try:
        data = request.json
        connection = DatabaseConfig.get_connection()
        cursor = connection.cursor()
        
        cursor.execute(
            "UPDATE categories SET name=%s, description=%s WHERE category_id=%s",
            (data.get('name'), data.get('description'), category_id)
        )
        updated_rows = cursor.rowcount
        
        connection.commit()
        cursor.close()
        connection.close()
        
        if updated_rows > 0:
            return jsonify({'success': True})
        else:
            return jsonify({'success': False, 'error': 'Category not found'}), 404
    except Exception as e:
        print(f"Error updating category {category_id}: {e}")
        try:
            if 'connection' in locals():
                connection.rollback()
        except Exception:
            pass
        import traceback
        traceback.print_exc()
        return jsonify({'success': False, 'error': f'Failed to update category: {str(e)}'}), 500

@app.route('/api/categories/<int:category_id>', methods=['DELETE'])
def delete_category(category_id):
    try:
        connection = DatabaseConfig.get_connection()
        cursor = connection.cursor()
        
        # The FK constraint is ON DELETE SET NULL, so files will be automatically
        # set to NULL when we delete the category. Just delete it directly.
        cursor.execute("DELETE FROM categories WHERE category_id = %s", (category_id,))
        deleted_rows = cursor.rowcount
        
        connection.commit()
        cursor.close()
        connection.close()
        
        if deleted_rows > 0:
            return jsonify({'success': True, 'message': 'Category deleted successfully'})
        else:
            return jsonify({'success': False, 'error': 'Category not found'}), 404
    except Exception as e:
        print(f"Error deleting category {category_id}: {e}")
        try:
            if 'connection' in locals():
                connection.rollback()
        except Exception:
            pass
        import traceback
        traceback.print_exc()
        return jsonify({'success': False, 'error': f'Failed to delete category: {str(e)}'}), 500



@app.route('/api/workspaces', methods=['GET'])
def get_workspaces():
    request_user = get_request_user()
    if not request_user:
        return jsonify({'error': 'Unauthorized'}), 401
    workspaces = DMSDatabase.get_all_workspaces() or []
    if request_user['role'] == 'system_admin':
        return jsonify(workspaces)
    
    visible = []
    for workspace in workspaces:
        members = workspace.get('members') or []
        if workspace.get('user_id') == request_user['user_id'] or request_user['user_id'] in members:
            visible.append(workspace)
    return jsonify(visible)

@app.route('/api/workspaces', methods=['POST'])
def create_workspace():
    try:
        data = request.json
        print(f"DEBUG: Creating workspace with data: {data}")
        
        workspace_id = DMSDatabase.create_workspace(
            data['name'], 
            data.get('description'), 
            data['user_id'], 
            data.get('members', [])
        )
        
        if workspace_id:
            return jsonify({'success': True, 'workspace_id': workspace_id})
        else:
            return jsonify({'error': 'Failed to create workspace'}), 500
            
    except Exception as e:
        print(f"Error in create_workspace API: {e}")
        return jsonify({'error': f'Failed to create workspace: {str(e)}'}), 500

@app.route('/api/settings', methods=['GET'])
def get_settings():
    settings = DMSDatabase.get_settings()
    return jsonify(serialize_settings_record(settings) or {})

@app.route('/api/settings', methods=['PUT'])
def update_settings():
    data = request.json
    DMSDatabase.update_settings(
        data['company'],
        data['max_file_mb'],
        data['allowed_types'],
        data.get('company_logo')
    )
    return jsonify({'success': True})


@app.route('/api/settings/logo', methods=['POST'])
def upload_company_logo():
    """
    Upload and save a global company logo image for the system.
    The stored filename is protected using the same helper that protects file paths.
    """
    try:
        if 'logo' not in request.files:
            return jsonify({'error': 'No logo file provided'}), 400

        file = request.files['logo']
        if file.filename == '':
            return jsonify({'error': 'Please select an image file'}), 400

        if not file.content_type.startswith('image/'):
            return jsonify({'error': 'File must be an image'}), 400

        # Basic size limit: 5MB
        file.seek(0, 2)
        size = file.tell()
        file.seek(0)
        if size > 5 * 1024 * 1024:
            return jsonify({'error': 'Logo size must be less than 5MB'}), 400

        upload_folder = os.path.join(os.path.dirname(__file__), 'uploads', 'company_logo')
        os.makedirs(upload_folder, exist_ok=True)

        import time
        file_ext = file.filename.rsplit('.', 1)[1].lower() if '.' in file.filename else 'png'
        raw_filename = f"logo_{int(time.time())}.{file_ext}"
        raw_path = os.path.join(upload_folder, raw_filename)

        file.save(raw_path)

        # Protect the stored filename using the same helper as other file paths
        protected_name = hash_file_path(raw_path)
        safe_token = f"{protected_name}.{file_ext}" if protected_name else raw_filename
        final_path = os.path.join(upload_folder, safe_token)

        try:
            if os.path.exists(final_path):
                os.remove(final_path)
            os.rename(raw_path, final_path)
            stored_token = safe_token
        except Exception as e:
            print(f"Company logo rename error: {e}")
            stored_token = raw_filename

        DMSDatabase.update_company_logo(stored_token)
        logo_url = build_company_logo_url(stored_token)

        return jsonify({'success': True, 'logo_url': logo_url, 'logo_token': stored_token})
    except Exception as e:
        print(f"Error uploading company logo: {e}")
        return jsonify({'error': f'Failed to upload company logo: {str(e)}'}), 500

@app.route('/api/activities', methods=['GET'])
def get_activities():
    request_user = get_request_user()
    if not request_user:
        return jsonify({'error': 'Unauthorized'}), 401
    
    # System admins see all activities, others see only their own
    if request_user.get('role') == 'system_admin':
        activities = DMSDatabase.get_recent_activities()
    else:
        activities = DMSDatabase.get_recent_activities_by_user(request_user['user_id'])
    return jsonify(activities or [])

@app.route('/api/activities', methods=['POST'])
def log_activity():
    data = request.json
    DMSDatabase.log_activity(data['user_id'], data['activity'])
    return jsonify({'success': True})

@app.route('/api/sessions', methods=['POST'])
def create_session():
    data = request.json
    session_id = DMSDatabase.create_session(data['user_id'], data['email'])
    return jsonify({'session_id': session_id})

@app.route('/api/reset-codes', methods=['POST'])
def create_reset_code():
    """
    Store reset code in database for password recovery.
    Expects: {email: str, code: str}
    """
    try:
        data = request.json
        email = data.get('email')
        code = data.get('code')
        
        if not email or not code:
            return jsonify({'success': False, 'error': 'Email and code are required'}), 400
        
        # Store reset code in database
        reset_code_id = DMSDatabase.create_reset_code(email, code)
        
        if reset_code_id:
            print(f"✅ Reset code stored in DB: email={email}, code={code}, id={reset_code_id}")
            return jsonify({'success': True, 'code': code, 'reset_code_id': reset_code_id}), 201
        else:
            print(f"❌ Failed to store reset code for {email}")
            return jsonify({'success': False, 'error': 'Failed to store reset code'}), 500
    except Exception as e:
        print(f"❌ Error in create_reset_code: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/reset-password', methods=['POST'])
def reset_password():
    """
    Reset user password using reset code.
    Marks the reset code as used after successful password update.
    Expects: {email: str, code: str, password: str}
    """
    try:
        data = request.json
        email = data.get('email')
        code = data.get('code')
        new_password = data.get('password')
        
        if not email or not code or not new_password:
            return jsonify({'success': False, 'error': 'Email, code, and password are required'}), 400
        
        # Verify the reset code is valid (not used and not expired)
        reset_record = DMSDatabase.verify_reset_code(email, code)
        if not reset_record:
            print(f"❌ Invalid or expired reset code for {email}")
            return jsonify({'success': False, 'error': 'Invalid or expired reset code'}), 400
        
        # Find user by email
        users = DMSDatabase.get_all_users()
        user = next((u for u in users if u['email'] == email), None)
        
        if not user:
            print(f"❌ User not found with email: {email}")
            return jsonify({'success': False, 'error': 'User not found'}), 404
        
        # Update user password
        update_result = DMSDatabase.update_user(
            user['user_id'],
            name=None,
            email=None,
            role=None,
            department_id=None,
            status=None,
            profile_image=None,
            password=new_password
        )
        
        if update_result:
            # Mark reset code as used
            DMSDatabase.mark_reset_code_as_used(email, code)
            print(f"✅ Password reset successful for {email}, reset code marked as used")
            return jsonify({'success': True, 'message': 'Password reset successfully'}), 200
        else:
            print(f"❌ Failed to update password for user {user['user_id']}")
            return jsonify({'success': False, 'error': 'Failed to reset password'}), 500
            
    except Exception as e:
        print(f"❌ Error in reset_password: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/users/<int:user_id>', methods=['GET'])
def get_user(user_id):
    user = DMSDatabase.get_user_by_id(user_id)
    if user:
        return jsonify(serialize_user_record(user))
    return jsonify({'error': 'User not found'}), 404

@app.route('/api/workspaces/<int:workspace_id>', methods=['GET'])
def get_workspace(workspace_id):
    request_user = get_request_user()
    if not request_user:
        return jsonify({'error': 'Unauthorized'}), 401
    workspaces = DMSDatabase.get_all_workspaces()
    workspace = next((w for w in workspaces if w['workspace_id'] == workspace_id), None)
    if not workspace:
        return jsonify({'error': 'Workspace not found'}), 404
    
    members = workspace.get('members') or []
    if request_user['role'] != 'system_admin' and request_user['user_id'] not in members and workspace.get('user_id') != request_user['user_id']:
        return jsonify({'error': 'Access denied'}), 403
    
    workspace['members'] = DMSDatabase.get_workspace_members(workspace_id)
    return jsonify(workspace)

@app.route('/api/workspaces/<int:workspace_id>', methods=['PUT'])
def update_workspace(workspace_id):
    try:
        data = request.json
        print(f"DEBUG: Updating workspace {workspace_id} with data: {data}")
        
        # Update workspace with members
        result = DMSDatabase.update_workspace(
            workspace_id, 
            data.get('name'), 
            data.get('description'),
            data.get('members')  # Pass members directly
        )
        
        if result:
            return jsonify({'success': True})
        else:
            return jsonify({'error': 'Failed to update workspace'}), 500
            
    except Exception as e:
        print(f"Error in update_workspace API: {e}")
        return jsonify({'error': f'Failed to update workspace: {str(e)}'}), 500
    
    
@app.route('/api/workspaces/<int:workspace_id>', methods=['DELETE'])
def delete_workspace(workspace_id):
    """Delete a workspace and all its associations"""
    try:
        # First, delete file-workspace associations
        delete_files_query = "DELETE FROM file_workspaces WHERE workspace_id = %s"
        DatabaseConfig.execute_query(delete_files_query, (workspace_id,))
        
        # Then delete workspace members
        delete_members_query = "DELETE FROM workspace_members WHERE workspace_id = %s"
        DatabaseConfig.execute_query(delete_members_query, (workspace_id,))
        
        # Finally delete the workspace itself
        delete_workspace_query = "DELETE FROM workspaces WHERE workspace_id = %s"
        DatabaseConfig.execute_query(delete_workspace_query, (workspace_id,))
        
        # If we reach here without exception, deletion was successful
        return jsonify({'success': True, 'message': 'Workspace deleted successfully'})
            
    except Exception as e:
        print(f"Error deleting workspace: {str(e)}")
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/api/files/<int:file_id>/add-to-workspace', methods=['POST'])
def add_file_to_workspace_api(file_id):
    """Add file to a workspace"""
    try:
        data = request.json
        workspace_id = data.get('workspace_id')
        
        if not workspace_id:
            return jsonify({'error': 'Workspace ID is required'}), 400
        
        result = DMSDatabase.add_file_to_workspace(file_id, workspace_id)
        
        if result:
            return jsonify({'success': True, 'message': 'File added to workspace'})
        else:
            return jsonify({'error': 'Failed to add file to workspace'}), 500
            
    except Exception as e:
        return jsonify({'error': f'Failed to add file to workspace: {str(e)}'}), 500

@app.route('/api/workspaces/<int:workspace_id>/files', methods=['GET'])
def get_workspace_files_api(workspace_id):
    """Get all files in a workspace using the new many-to-many relationship"""
    try:
        request_user = get_request_user()
        if not request_user:
            return jsonify({'error': 'Unauthorized'}), 401
        
        workspaces = DMSDatabase.get_all_workspaces() or []
        workspace = next((w for w in workspaces if w['workspace_id'] == workspace_id), None)
        if not workspace:
            return jsonify({'error': 'Workspace not found'}), 404
        
        members = workspace.get('members') or []
        if request_user['role'] != 'system_admin' and request_user['user_id'] not in members and workspace.get('user_id') != request_user['user_id']:
            return jsonify({'error': 'Access denied'}), 403
        
        print(f"🚀 DEBUG: API - Getting files for workspace {workspace_id}")
        files = DMSDatabase.get_files_by_workspace(workspace_id)
        return jsonify(files)
    except Exception as e:
        print(f"❌ Error getting workspace files: {e}")
        return jsonify({'error': f'Failed to get workspace files: {str(e)}'}), 500
    
    
@app.route('/api/users/<int:user_id>', methods=['DELETE'])
def delete_user(user_id):
    """Permanently delete a user and handle all related records"""
    try:
        # First check if user exists
        user = DMSDatabase.get_user_by_id(user_id)
        if not user:
            return jsonify({'error': 'User not found'}), 404
            
        connection = DatabaseConfig.get_connection()
        cursor = connection.cursor()
        
        try:
            # Delete in order to respect foreign key constraints
            
            # 1. Delete user's file shares (as shared_by or shared_with)
            cursor.execute("DELETE FROM file_shares WHERE shared_by = %s OR shared_with_user_id = %s", (user_id, user_id))
            print(f"Deleted file shares for user {user_id}")
            
            # 2. Delete user's workspace memberships
            cursor.execute("DELETE FROM workspace_members WHERE user_id = %s", (user_id,))
            print(f"Deleted workspace memberships for user {user_id}")
            
            # 3. Delete user's sessions
            cursor.execute("DELETE FROM sessions WHERE user_id = %s", (user_id,))
            print(f"Deleted sessions for user {user_id}")
            
            # 4. Delete user's activity logs
            cursor.execute("DELETE FROM activity_log WHERE user_id = %s", (user_id,))
            print(f"Deleted activity logs for user {user_id}")
            
            # 5. Delete or reassign user's files (delete them since user is being removed)
            cursor.execute("DELETE FROM files WHERE user_id = %s", (user_id,))
            print(f"Deleted files for user {user_id}")
            
            # 6. Delete or reassign user's workspaces (delete them)
            cursor.execute("DELETE FROM workspaces WHERE user_id = %s", (user_id,))
            print(f"Deleted workspaces for user {user_id}")
            
            # 7. Finally, delete the user
            cursor.execute("DELETE FROM users WHERE user_id = %s", (user_id,))
            print(f"Deleted user {user_id}")
            
            connection.commit()
            cursor.close()
            connection.close()
            
            return jsonify({'success': True, 'message': 'User permanently deleted'})
            
        except Exception as e:
            connection.rollback()
            cursor.close()
            connection.close()
            print(f"Error during user deletion: {e}")
            import traceback
            traceback.print_exc()
            return jsonify({'error': f'Database error: {str(e)}'}), 500
            
    except Exception as e:
        print(f"Error deleting user: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({'error': f'Failed to delete user: {str(e)}'}), 500

@app.route('/api/files/<int:file_id>/content', methods=['GET'])
def get_file_content(file_id):
    """Get file content as text for editing"""
    try:
        file_record = DMSDatabase.get_file_by_id(file_id)
        if not file_record:
            return jsonify({'error': 'File not found'}), 404
        
        request_user = get_request_user()
        if not request_user:
            return jsonify({'error': 'Unauthorized'}), 401
        
        permission = get_user_file_permission(file_record, request_user)
        if not permission:
            return jsonify({'error': 'Access denied'}), 403
        
        actual_path = resolve_file_path(file_record)
        if not actual_path or not os.path.exists(actual_path):
            return jsonify({'error': 'File not found on server'}), 404
        
        try:
            content = extract_text_for_editing(file_record, actual_path)
            return jsonify({'success': True, 'content': content})
        except ValueError as ve:
            return jsonify({'error': str(ve)}), 400
        except Exception as e:
            return jsonify({'error': f'Failed to read file: {str(e)}'}), 500
            
    except Exception as e:
        return jsonify({'error': f'Failed to get file content: {str(e)}'}), 500

@app.route('/api/files/<int:file_id>/edit', methods=['POST'])
def save_file_edit(file_id):
    """Save edited file content"""
    try:
        data = request.json
        content = data.get('content')
        
        if content is None:
            return jsonify({'error': 'Content is required'}), 400
        
        file_record = DMSDatabase.get_file_by_id(file_id)
        if not file_record:
            return jsonify({'error': 'File not found'}), 404
        
        request_user = get_request_user()
        if not request_user:
            return jsonify({'error': 'Unauthorized'}), 401
        
        permission = get_user_file_permission(file_record, request_user)
        if permission not in ('owner', 'editor', 'admin'):
            return jsonify({'error': 'You do not have permission to edit this file.'}), 403
        
        actual_path = resolve_file_path(file_record)
        if not actual_path or not os.path.exists(actual_path):
            return jsonify({'error': 'File not found on server'}), 404
        
        try:
            save_text_back_to_file(file_record, actual_path, content)
            DMSDatabase.update_file(file_id, {'updated_at': datetime.now()})
            return jsonify({'success': True, 'message': 'File saved successfully'})
        except ValueError as ve:
            return jsonify({'error': str(ve)}), 400
        except Exception as e:
            return jsonify({'error': f'Failed to save file: {str(e)}'}), 500
            
    except Exception as e:
        return jsonify({'error': f'Failed to save file edit: {str(e)}'}), 500

@app.route('/api/health', methods=['GET'])
def health_check():
    return jsonify({'status': 'healthy', 'service': 'document-classifier'})

if __name__ == '__main__':
    app.run(debug=True, port=5000)




